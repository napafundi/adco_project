from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkinter import filedialog
import sqlite3
from tkcalendar import Calendar
import os
import webbrowser
from PIL import Image
from PIL import ImageTk
import re
from datetime import datetime
from datetime import date
import calendar
import math
import openpyxl
from openpyxl.styles import Font
import string
from docx import Document
from docx.shared import Pt
import collections
import pandas
import matplotlib
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg


def database():
    #Create inventory database if it does not exist.  Update certain
    #values within the database.
    conn = sqlite3.Connection("inventory.db")
    cur = conn.cursor()
    cur.execute("""CREATE TABLE IF NOT EXISTS 'raw_materials'
                (type TEXT, product TEXT, amount INTEGER, price REAL,total TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'production'
                (date DATE, product TEXT, amount INTEGER)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'in_progress'
                (date DATE, product TEXT, amount INTEGER, description INTEGER)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'bottles'
                (type TEXT, product TEXT, amount INTEGER,
                 case_size INTEGER, price REAL, total TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'samples'
                (type TEXT, product TEXT, amount INTEGER,
                 price REAL, total TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'grain'
                (date DATE, 'order_number' TEXT, type TEXT, amount INTEGER,
                 price REAL, total TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'grain_log'
                (arrival_date DATE, finish_date DATE, type TEXT, order_no TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'mashes'
                (date DATE, mash_no TEXT, type TEXT, grains TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'barrels'
                ('barrel_no' TEXT, type TEXT, gallons INTEGER,
                 'proof_gallons' REAL, 'date_filled' DATE, age TEXT,
                 investor TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'empty_barrels'
                ('barrel_no' TEXT, type TEXT, gallons INTEGER,
                 'proof_gallons' REAL,'pg_remaining' REAL, 'date_filled' DATE,
                 'date_emptied' DATE, age TEXT, investor TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'barrel_count'
                ('full_amount' INTEGER,'empty_amount' INTEGER, price REAL,
                 total TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'estimated_cogs'
                (raw_mat REAL, energy REAL, labor REAL, error REAL,
                 'total_per_bottle' REAL, 'bond_ins' REAL, storage REAL,
                 mult_fact REAL, 'total_per_pg' REAL)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'purchase_orders'
                (po_date DATE, pu_date DATE, product TEXT, amount INTEGER,
                 unit TEXT, price REAL, total REAL, destination TEXT,
                 'po_number' TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'pending_po'
                (po_date DATE, pu_date DATE, product TEXT, amount INTEGER,
                 unit TEXT, price REAL, total REAL, destination TEXT,
                 'po_number' TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'employee_transactions'
                (date DATE, product TEXT, amount INTEGER, unit TEXT,
                 employee TEXT, destination TEXT)
                """)
    cur.execute("""CREATE TABLE IF NOT EXISTS 'monthly_reports'
                (date DATE, inv_name TEXT, total REAL,
                 UNIQUE(date, inv_name) ON CONFLICT REPLACE)
                """)
    conn.commit()
    conn.close()
    db_update()


def db_update():
    #Updates inventory database values for specific columns.
    conn = sqlite3.Connection("inventory.db")
    cur = conn.cursor()
    cur.execute("""UPDATE 'raw_materials'
                      SET total=PRINTF('%s%g', '$', amount*price)
                """)
    cur.execute("""UPDATE 'raw_materials'
                      SET price=PRINTF('%.2f', price)
                """)
    cur.execute("""UPDATE 'bottles'
                      SET amount=0
                    WHERE amount<0
                """)
    cur.execute("""UPDATE 'bottles'
                      SET total=PRINTF('%s%g', '$', amount*price)
                """)
    cur.execute("""UPDATE 'samples'
                      SET amount=0
                    WHERE amount<0
                """)
    cur.execute("""UPDATE 'samples'
                      SET total=PRINTF('%s%.2f', '$', amount*price)
                """)
    cur.execute("""UPDATE 'grain'
                      SET total=PRINTF('%s%g', '$', amount*price)
                """)
    cur.execute("""UPDATE 'barrels'
                      SET age=PRINTF('%d years, %d months',
                                     ((julianday('now')
                                      - julianday(date_filled)) / 365),
                                     (((julianday('now')
                                      - julianday(date_filled)) % 365) / 30))
                """)
    cur.execute("""SELECT COUNT(*)
                     FROM 'barrels'
                """)
    barrel_count = str(cur.fetchone()[0])
    cur.execute("UPDATE 'barrel_count' " +
                   "SET full_amount=" + barrel_count)
    cur.execute("""UPDATE 'barrel_count'
                      SET total=printf('%s%.2f', '$',
                                       (full_amount + empty_amount) * price)
                """)
    cur.execute("""UPDATE 'estimated_cogs'
                      SET total_per_bottle=PRINTF('%.2f',
                                                  (raw_mat + energy + labor
                                                   + error))
                """)
    cur.execute("""UPDATE 'estimated_cogs'
                    SET total_per_pg=PRINTF('%.2f',
                                            ((total_per_bottle*mult_fact)
                                             + bond_ins + storage))
                """)
    conn.commit()
    conn.close()

def monthly_reports_update():
    #Update 'monthly_reports' table with current month and inventory
    #values. Only select purchase orders from current month. Selects
    #pending purchase orders to be fulfilled next month.
    inv_tables = ['raw_materials', 'bottles', 'samples', 'pending_po', 'grain',
                  'purchase_orders']
    monthly_totals = collections.OrderedDict()
    cur_date = datetime.today().strftime('%Y-%m')
    conn = sqlite3.Connection("inventory.db")
    cur = conn.cursor()
    cur.execute("SELECT product, price FROM bottles")
    prod_prices = {key:val for (key, val) in cur.fetchall()}
    for table in inv_tables:
        if table == 'pending_po':
            cur.execute("SELECT product, amount, total " +
                          "FROM " + table +
                        " WHERE po_date " +
                          "LIKE \'" + cur_date + "%\'")
            pending_info = cur.fetchall()
            pend_sale_amts = {} #Total sale amounts by product
            for (prod, amt) in [x[:2] for x in pending_info]:
                if prod in pend_sale_amts:
                    pend_sale_amts[prod] += int(amt)
                else:
                    pend_sale_amts[prod] = int(amt)
            pend_cogs_total = 0
            for prod in pend_sale_amts.keys():
                pend_cogs_total += float(pend_sale_amts[prod]
                                         * prod_prices[prod])
            pend_sale_total = sum([float(x[2][1:].replace(",", "")) for x
                                   in pending_info])
            monthly_totals['pending_sales'] = pend_sale_total
            monthly_totals['pending_cogs'] = -1*pend_cogs_total
        elif table == 'purchase_orders':
            cur.execute("SELECT product, amount, total " +
                          "FROM " + table +
                        " WHERE pu_date " +
                          "LIKE \'" + cur_date + "%\'")
            po_info = cur.fetchall()
            po_sale_amts = {}   #Total sales amount by product
            for (prod, amt) in [x[:2] for x in po_info]:
                if prod in po_sale_amts.keys():
                    po_sale_amts[prod] += int(amt)
                else:
                    po_sale_amts[prod] = int(amt)
            try:
                po_cogs_total = 0
                for prod in po_sale_amts.keys():
                    po_cogs_total += float(po_sale_amts[prod]
                                           * prod_prices[prod])
            except:
                po_cogs_total = 0
            po_sale_total = sum([float(x[2][1:].replace(",", "")) for x
                                 in po_info])
            monthly_totals['purchase_order_sales'] = po_sale_total
            monthly_totals['purchase_order_cogs'] = -1*po_cogs_total
        else:
            cur.execute("SELECT total " +
                          "FROM " + table)
            total_vals = [x[0] for x in cur.fetchall()]
            total = sum([float(x[1:].replace(",", "")) for x
                         in total_vals])
            monthly_totals[table] = total
    cur.execute("SELECT * " +
                  "FROM barrels")
    barrel_vals = cur.fetchall()
    cur.execute("""SELECT total_per_pg
                     FROM estimated_cogs
                """)
    est_cogs = [x[0] for x in cur.fetchall()]
    whisk_cogs = float(est_cogs[0])
    rum_cogs = float(est_cogs[1])
    whisk_total = 0
    rum_total = 0
    for barrel in barrel_vals:
        if barrel[1] == 'Rum':
            rum_total += float(barrel[3]) * rum_cogs
        else:
            whisk_total += float(barrel[3]) * whisk_cogs
    monthly_totals['barreled_rum'] = ("%.2f" % rum_total)
    monthly_totals['barreled_whiskey'] = ("%.2f" % whisk_total)
    cur.execute("""SELECT total
                     FROM barrel_count
                """)
    barr_total = cur.fetchone()[0][1:].replace(",", "")
    monthly_totals['barrels'] = barr_total
    for key, value in monthly_totals.items():
        value = ("%.2f" % float(value))
        cur.execute("INSERT INTO monthly_reports " +
                         "VALUES (?, ?, ?)",
                         (cur_date, key, value))
    conn.commit()
    conn.close()


def create_excel_inv():
    #Populate inventory_template.xlsx with inventory values and save as
    #new workbook.
    #('inv table', 'total column index')
    inventories = (('raw_materials', 4),
                   ('production', None),
                   ('in_progress', None),
                   ('bottles', 5),
                   ('samples', 5),
                   ('grain', 5),
                   ('mashes', None),
                   ('grain_log', None),
                   ('barrels', None),
                   ('empty_barrels', None),
                   ('purchase_orders', 6),
                   ('pending_po', 6),
                   ('employee_transactions', None))
    inventories = collections.OrderedDict(inventories)
    excel_file = (os.getcwd() + "/inventory_files/inventory_template.xlsx")
    wb = openpyxl.load_workbook(excel_file)
    sheets = wb.sheetnames
    conn = sqlite3.Connection("inventory.db")
    cur = conn.cursor()
    for sheet, (inv, tot_col) in zip(sheets, inventories.items()):
        active_sheet = wb[sheet]
        cur.execute("SELECT * " +
                      "FROM " + inv)
        inv_values = cur.fetchall()
        if len(inv_values) > 0:
            for (indx, row) in enumerate(inv_values, 2):
                try:
                    row = list(row)
                    # Strip $, remove commas
                    row[tot_col] = float(row[tot_col][1:].replace(",",""))
                    tot_col += 1    # Change index for excel column
                    active_sheet.append(row)
                    active_sheet.cell(indx, tot_col).number_format = '$#,##0.00'
                except TypeError:
                    pass
            last_row = len(inv_values) + 1
            last_col = len(inv_values[0]) - 1
            # Format text to justify centrally
            rows = range(1, last_row + 1)
            columns = string.ascii_uppercase[:last_col + 1]
            last_col = string.ascii_uppercase[last_col]
            for row in rows:
                for col in columns:
                    cell = col + str(row)
                    active_sheet[cell].alignment = (
                        openpyxl.styles.Alignment(horizontal='center'))
            # Format excel table
            tbl_ref = "A1:" + str(last_col) + str(last_row)
            tbl = openpyxl.worksheet.table.Table(displayName=inv, ref=tbl_ref)
            style = openpyxl.worksheet.table.TableStyleInfo(
                name="TableStyleMedium9", showFirstColumn=False,
                showLastColumn=False, showRowStripes=True)
            tbl.tableStyleInfo = style
            active_sheet.add_table(tbl)
    conn.close()
    new_excel_file = (os.getcwd() + "/inventory_files/"
                      + datetime.now().strftime("%Y-%m") + ".xlsx")
    wb.save(new_excel_file)
    os.system('start EXCEL.EXE ' + new_excel_file)


def edit_db(sql_edit, sqlite_table, gui_table, view_fr, delete=False):
    #Updates the sqlite_table with the changes provided by sql_edit.
    #sql_edit is a tuple of length 2*(num of cols)
    conn = sqlite3.Connection("inventory.db")
    cur = conn.cursor()
    cur.execute("SELECT * " +
                  "FROM " + sqlite_table)
    columns = [x[0] for x
               in cur.description]
    str1 = "=?, ".join(columns) + "=?"  # row1=?, row2=?, ...
    str2 = "=? AND ".join(columns) + "=?"   # row1=? AND row2=? AND ...
    if delete == False:
        cur.execute("UPDATE " + sqlite_table +
                      " SET " + str1 +
                    " WHERE " + str2,
                    sql_edit)
    else:
        cur.execute("DELETE FROM " + sqlite_table +
                         " WHERE " + str2,
                    sql_edit)
    conn.commit()
    conn.close()
    try:
        view_fr.columns.set("All")
        view_fr.columns.event_generate("<<ComboboxSelected>>")
    except:
        pass
    try:
        view_products(sqlite_table, 'All', 'All', gui_table)
    except:
        pass
    db_update()
    if sqlite_table == 'barrels':
        barr_count_fr.barr_update(first=1)


def view_widget(window, widget, location, sqlite_table, column, item,
                gui_table):
    #Removes current packed widgets from window frame and replaces with
    #new widget chosen.
    for widg in window.pack_slaves():
        widg.pack_forget()
    widget.pack(side=location, fill=BOTH, expand=1)
    if gui_table:
        view_products(sqlite_table, column, item, gui_table)


def view_products(sqlite_table, column, item, gui_table):
    #Fetches info from sqlite_table based on an item filter.  Returns
    #information into the current gui_table.  Configures even-numbered
    #rows to have a grey background.
    conn = sqlite3.Connection("inventory.db")
    cur = conn.cursor()
    if column == "All":
        cur.execute("SELECT * " +
                      "FROM " + sqlite_table)
    elif column == "barrel_no":
        cur.execute("SELECT * " +
                      "FROM " + sqlite_table +
                    " WHERE " + column +
                     " LIKE \'" + item[:2] + "%\'")
    elif 'date' in column:
        try:
            cur.execute("SELECT * " +
                          "FROM " + sqlite_table +
                        " WHERE " + column +
                         " LIKE \'" + item[:4] + "%\'")
        except sqlite3.OperationalError:
            cur.execute("SELECT * " +
                          "FROM " + sqlite_table +
                        " WHERE po_date" +
                         " LIKE \'" + item[:4] + "%\'")
    elif 'pick' in column:
        cur.execute("SELECT * " +
                      "FROM " + sqlite_table +
                    " WHERE pu_date" +
                     " LIKE \'" + item[:4] + "%\'")
    elif column == "age":
        cur.execute("SELECT * " +
                      "FROM " + sqlite_table +
                    " WHERE " + column +
                     " LIKE \'" + item[0] + "%\'")
    elif column == "product":
        cur.execute("SELECT * " +
                      "FROM " + sqlite_table +
                    " WHERE " + column +
                     " LIKE \'" + item + "%\'")
    else:
        cur.execute("SELECT * " +
                      "FROM " + sqlite_table +
                    " WHERE " + column + "= \'" + item + "\'")
    rows = cur.fetchall()
    conn.close()

    for item in gui_table.get_children():
        gui_table.delete(item)

    for index,row in enumerate(rows):
        if (index % 2 == 0):
            tag = 'even'
        else:
            tag = 'odd'
        gui_table.insert("", END, values=row, tags=(tag,))
    gui_table.tag_configure('even', background='#E8E8E8')


def file_view(folder):
    #Displays a toplevel window populated by clickable links to
    #files within the given folder.
    labels_window = Toplevel(window)
    files = os.listdir(os.getcwd() + "\\" + folder)
    window_height = 0
    for file in files:
        mo = fileRegex.search(file)
        file_name = mo.group(1).replace("_", " ")
        file_label = Sheet_Label(
            master=labels_window, text=file_name,
            file_location=(os.getcwd() + "\\" + folder + "\\" + file))
        file_label.pack(padx=10, pady=5, anchor='w')
        window_height += 38
    labels_window.title(folder.replace("_", " ").title())
    labels_window.focus()
    x = (screen_width/2) - (250)
    y = (screen_height/2) - (250)
    labels_window.geometry("%dx%d+%d+%d" % (300, window_height, x, y))
    labels_window.resizable(0,0)


def selection_check(sqlite_table, gui_table, view_fr, edit=True, delete=False,
                    empty=False):
    #Checks to see if a gui_table selection has been made and returns
    #the respective action based on the gui_table.
    item_values = gui_table.item(gui_table.selection())['values']
    if item_values:
        if delete == True:
            del_ques = messagebox.askquestion(
            "Delete Current Selection?",
            "Are you sure you want to continue? Confirming will delete the " +
            "current selection from the inventory. This information will not " +
            "be able to be recovered.")
            if del_ques == 'yes':
                edit_db(tuple(item_values), sqlite_table, gui_table, view_fr,
                        delete=True)
            else:
                return
        elif empty == True:
            Empty_Barrel_View(window, item_values)
        elif (gui_table == po_tbl and edit==False): # Open po excel file
            po_num = item_values[8]
            try:
                excel_file = (os.getcwd() + "/purchase_orders/" + po_num[:4]
                              + "/" + po_num + ".xlsx")
                os.system('start EXCEL.EXE ' + excel_file)
            except:
                messagebox.showerror(
                    "Program Error",
                    "There was an error opening Excel.", parent=window)
        elif gui_table == inprog_tbl:   #Finish in progress production
            Finish_View(window, item_values)
        elif (gui_table == pending_tbl and edit==False):  #fulfill po
            fulfill_pending(gui_table, view_fr)
        else:   #edit selection
            Edit_View(window, sqlite_table, gui_table, 2, view_fr)
    else:
        messagebox.showerror(
            "Selection Error",
            "Please select an inventory item.", parent=window)


def gui_table_sort(gui_table, column, reverse):
    #Sorts gui tables in ascending order based on the column header
    #clicked.  The next click upon the header will be in reverse order.
    l = [(gui_table.set(k, column), k) for k
         in gui_table.get_children()]
    if '$' in l[0][0]:  #Check if column is 'total'
        l.sort(key=lambda tup: float(tup[0][1:].replace(",", "")),
               reverse=reverse)
    else:
        try:
            l.sort(key=lambda tup: float(tup[0].replace(",","")),
                   reverse=reverse)
        except ValueError:
            l.sort(key=lambda tup: tup[0], reverse=reverse)

    #Rearrange items in sorted positions.
    for index, (val, k) in enumerate(l):
        gui_table.move(k, '', index)
        gui_table.item(k, tags=())
        if index % 2 == 0:
            gui_table.item(k, tags=('even',))
    gui_table.tag_configure('even', background="#E8E8E8")

    #Reverse sort next time.
    gui_table.heading(
        column, text=column,
        command=lambda c=column: gui_table_sort(gui_table, c, not reverse))


def cal_button(tplvl, date_entry):
    #Creates a toplevel window to provide a calendar date selection
    #tool.
    tplvl.top = Toplevel(window)
    tplvl.cal = Calendar(tplvl.top, font="Arial 14", selectmode='day',
                         locale='en_US', cursor="hand2")
    tplvl.cal.pack(fill="both", expand=True)
    (HoverButton(tplvl.top, text="ok",
                 command=lambda: retrieve_date(tplvl, date_entry))
                 .pack())
    tplvl.top.focus()


def retrieve_date(tplvl, date_entry):
    #Updates the date-entry widget within the toplevel widget with
    #formatted date value.
    date_entry.config(state=NORMAL)
    date_entry.delete(0 ,END)
    date_entry.insert(END, tplvl.cal.selection_get().strftime("%Y-%m-%d"))
    date_entry.config(state="readonly")
    tplvl.top.destroy()


def confirm_po(view_fr, info, purchase_orders, po_num):
    #Insert purchase order info into 'purchase_orders' table and create
    #excel file with purchase order info.
    year = datetime.now().year
    wb = openpyxl.load_workbook(
             'purchase_orders/blank_po.xlsx')
    sheet = wb['Purchase Order']
    font = Font(name='Times New Roman', size=12)

    #Get shipment information entry-values into list and input them into
    #corresponding cells within the 'po' excel sheet.
    info_cells = ['A9', 'K9', 'A12', 'A15', 'I15']
    for entry,cell in zip(info, info_cells):
        sheet[cell] = entry
        sheet[cell].font = font
    #Get order values and input into table within purchase order excel
    #file.
    excel_rows = ["A","B","D","J","M"]
    excel_columns = range(18, 36)
    index = 0
    total_po = 0
    for i in excel_columns:
        for j,k in zip(excel_rows, range(0,5)):
            cell = j + str(i)
            try:
                sheet[cell] = (purchase_orders[index][k])
                sheet[cell].font = font
                if k == 4:
                    try:
                        total_po += float(
                            purchase_orders[index][k][1:].replace(',',''))
                    except ValueError:
                        pass
            except IndexError:
                sheet[cell] = ""
        index += 1
    total_po = "{0:,.2f}".format(total_po)
    sheet['M36'] = ("$%s" % total_po)

    #Add purchase orders to 'purchase_orders' table and update
    #inventory.
    conn = sqlite3.Connection("inventory.db")
    cur = conn.cursor()
    for po_list in (x for x in purchase_orders if all(x)):
        cur.execute("INSERT INTO 'purchase_orders' " +
                              "VALUES (?,?,?,?,?,?,?,?,?)",
                         (info[3], info[4], po_list[2], po_list[0], po_list[1],
                          po_list[3], po_list[4], info[2], po_num))
        if po_list[1] == "Cases":
            cur.execute("UPDATE 'bottles' " +
                           "SET amount=(amount - ?) " +
                         "WHERE product=?",
                        (po_list[0], po_list[2]))
        else:
            cur.execute("UPDATE 'samples' " +
                                "SET amount=(amount - ?) " +
                              "WHERE product=?",
                             (po_list[0], po_list[2]))
    conn.commit()
    conn.close()
    excel_file = (os.getcwd() + "/purchase_orders/"
                       + str(year) + "/" + po_num
                       + ".xlsx")
    wb.save(excel_file)

    open_ques = messagebox.askquestion(
        "Open the PO Excel File?",
        "Would you like to open the Purchase Order file in Excel? "
        + "This will allow you to print it now.")
    if open_ques == "yes":
        try:
            os.system('start EXCEL.EXE ' + excel_file)
        except:
            messagebox.showerror(
                "Program Error",
                "There was an error opening Excel.", parent=self)
    else:
        pass

    db_update()
    view_fr.columns.set("All")
    view_fr.columns.event_generate("<<ComboboxSelected>>")

def fulfill_pending(gui_table, view_fr):
    #Remove pending purchase order from 'pending_po' table and input
    #into 'purchase_orders'. Create an excel file with info.
    po_num = gui_table.item(gui_table.selection())['values'][8]
    conn = sqlite3.Connection("inventory.db")
    cur = conn.cursor()
    cur.execute("SELECT * " +
                  "FROM 'pending_po' " +
                 "WHERE po_number=\'" + po_num + "\'")
    po_vals = cur.fetchall()
    conn.close()
    po_1 = po_vals[0]
    po_info = ["Montgomery", po_num, po_1[7], po_1[0], po_1[1]]
    comp_pos = [[x[3] ,x[4], x[2], x[5], x[6]] for x in po_vals]
    confirm_po(view_fr, po_info, comp_pos, po_num)
    for po in po_vals:
        edit_db(po, 'pending_po', gui_table, view_fr, delete=True)

class HoverButton(Button):
    #Button widget with mouse-over color and cursor changes.
    def __init__(self, master, **kw):
        Button.__init__(self, master=master, cursor="hand2", **kw)
        self.defaultBackground = self["background"]
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)

    def on_enter(self, event):
        self['background'] = 'gray70'

    def on_leave(self, event):
        self['background'] = self.defaultBackground


class Add_View(Toplevel):
    #A toplevel widget with labels corresponding to sqlite table columns
    #and entry widgets to insert data into the sqlite table.
    def __init__(self, master, sqlite_table, gui_table, entry_col, view_fr):
        self.sqlite_table = sqlite_table
        self.gui_table = gui_table
        self.entry_col = entry_col
        self.view_fr = view_fr
        self.tplvl_title = self.sqlite_table.replace("_"," ").title()
        self.x = (screen_width/2) - (width/2) + 100
        self.y = ((screen_height/2) - (height/2)) + 50
        Toplevel.__init__(self, master=master)

        self.title_frame = Frame(self)
        (Label(self.title_frame,
               text="Add Product to " + self.tplvl_title + " Inventory",
               font="Arial 10 bold")
              .pack())
        self.title_frame.grid(row=0, column=0, columnspan=2, pady=5)

        #Create labels and entries based on gui_table column headers.
        for index,description in enumerate(self.gui_table.columns, 1):
            if (description.lower() != 'type'):
                Label(self, text=description + ":").grid(row=index, column=0)
                if description.lower() == 'total':
                    self.total_text = StringVar()
                    self.total_entry = Entry(self, textvariable=self.total_text)
                    self.total_entry.config(state="readonly")
                    self.total_entry.grid(row=index, column=self.entry_col)
                elif (description.lower().find('age') != -1):
                    self.age_text = StringVar()
                    self.age_entry = Entry(self, textvariable=self.age_text)
                    self.age_entry.config(state="readonly")
                    self.age_entry.grid(row=index, column=self.entry_col)
                else:
                    Entry(self).grid(row=index,column=self.entry_col)
                if (description.lower().find('date') != -1):
                    self.date_index = index
                    self.date_entry = self.grid_slaves(row=self.date_index,
                                                       column=self.entry_col)[0]
                    self.date_entry.config(state="readonly")
                    self.cal_link = HoverButton(
                        self, image=cal_photo,
                        command=lambda: cal_button(self, self.date_entry))
                    self.cal_link.image = cal_photo
                    self.cal_link.grid(row=index, column=self.entry_col+1)
                elif ((description.lower().find('total') != -1) or
                      (description.lower().find('age') != -1)):
                    self.labels = [x for x
                                   in reversed(self.grid_slaves(column=0))
                                   if (x.winfo_class() == 'Label')]
                    for entry in self.labels:
                        if entry.cget("text").lower().find("amount") != -1:
                            self.amount_row = entry.grid_info()['row']
                            self.amount_entry = self.grid_slaves(
                                row=self.amount_row,
                                column=self.entry_col)[0]
                        if entry.cget("text").lower().find("price") != -1:
                            self.price_row = entry.grid_info()['row']
                            self.price_entry = self.grid_slaves(
                                row=self.price_row,
                                column=self.entry_col)[0]
                    self.total_after()
            else:   #Type option entry.
                Label(self, text=description + ":").grid(row=index, column=0)
                self.options = ttk.Combobox(self,
                                            values=type_options[sqlite_table])
                self.options.set(type_options[sqlite_table][0])
                self.options.config(width=16, background="white",
                                    justify='center', state='readonly')
                self.options.grid(row=index, column=self.entry_col)

        self.grid_size = self.grid_size()[1]
        self.button_frame = Frame(self)
        (HoverButton(self.button_frame, text="Add Item", width=10,
                command=self.add_item)
                .pack(side=LEFT, padx=5, pady=5))
        (HoverButton(self.button_frame, text="Cancel", width=10,
                command=lambda : self.destroy())
                .pack(side=LEFT, padx=5, pady=5))
        self.button_frame.grid(row=self.grid_size+1, column=0, columnspan=2)

        self.title("Add to " + self.tplvl_title)
        self.focus()
        self.geometry("+%d+%d" % (self.x, self.y))
        self.resizable(0,0)

    def add_item(self):
        #Work through Add_View toplevel to find entry widgets and
        #extract these values to be inserted into the sqlite table.
        #Uses db_update() to update certain column values afterwards
        #and view_products() to display the updated gui table.
        self.additions = [] #list to be populated by entries
        self.num_entries = 0
        self.entries = [x for x
                        in reversed(self.grid_slaves())
                        if (x.winfo_class() == 'Entry'
                        or x.winfo_class() == 'TCombobox')]
        for entry in self.entries:
            if entry.winfo_class() == 'Entry' and entry.get():
                #Titlecase before appending
                self.additions.append(' '.join(word[0].upper() + word[1:]
                                      for word
                                      in entry.get().split()))
                self.num_entries += 1
            elif entry.winfo_class() == 'TCombobox':
                self.additions.append(entry.get())
                self.num_entries += 1
            else:
                messagebox.showerror("Input Error",
                                     "At least one input is blank, "
                                     + "please try again.", parent=self)
                return
        self.additions = tuple(self.additions)

        self.str1 = "?,"*(self.num_entries - 1) + "?)"  #?,?,..,?)
        self.conn = sqlite3.Connection("inventory.db")
        self.cur = self.conn.cursor()
        self.cur.execute("INSERT INTO " + self.sqlite_table +
                             " VALUES (" + self.str1,
                         self.additions)
        self.conn.commit()
        self.conn.close()
        db_update()
        try:
            self.view_fr.columns.set("All")
            self.view_fr.columns.event_generate("<<ComboboxSelected>>")
        except:
            view_products(self.sqlite_table, 'All', 'All', self.gui_table)
        if self.sqlite_table == 'barrels':
            barr_count_fr.barr_update(first=1)
        self.destroy()

    def total_after(self):
        #Widget after-function to update total and age entry values.
        #Repeats every 150ms.
        def total_update():
            #Tries to update total and age entry values.
            #Raises:
            #AttributeError: if price_entry, amount_entry, date_entry
            #don't exist
            #ValueError: if price_entry, amount_entry, date_entry values
            #are currently empty
            try:    #update total entry
                self.price_num = self.price_entry.get()
                self.amount_num = self.amount_entry.get()
                self.total_string = ("$%.2f" %
                                     (float(self.amount_num)
                                     *float(self.price_num)))
                self.total_text.set(self.total_string)
                return
            except (AttributeError, ValueError):
                pass
            try:    #Update age entry.
                self.date_value = datetime.strptime(self.date_entry.get(),
                                                    '%Y-%m-%d')
                self.date_diff = datetime.now() - self.date_value
                self.barrel_age = ("%d years, %d months" %
                                   (math.floor(self.date_diff.days/365.2425),
                                   (self.date_diff.days%365.2425)/30))
                self.age_text.set(self.barrel_age)
            except (AttributeError, ValueError):
                pass
        total_update()
        self.after(150, self.total_after)


class Edit_View(Add_View):
    #A toplevel widget with labels corresponding to sqlite table columns
    #and entry widgets to update data in sqlite table.
    def __init__(self, master, sqlite_table, gui_table, entry_col, view_fr):
        self.master = master
        self.sqlite_table = sqlite_table
        self.gui_table = gui_table
        self.entry_col = entry_col  #Column location for entry widgets.
        self.view_fr = view_fr
        self.selection = self.gui_table.selection()
        self.item_values = self.gui_table.item(self.selection)['values']
        self.tplvl_title = self.sqlite_table.replace("_"," ").title()
        Add_View.__init__(self, master, sqlite_table, gui_table, entry_col,
                          view_fr)
        self.title("Edit " + self.tplvl_title)
        self.title_frame.destroy()  #Remove add_view title frame.

        self.title_frame = Frame(self)
        (Label(self.title_frame,
               text="Edit Product in " + self.tplvl_title + " Inventory",
               font="Arial 10 bold")
               .pack())
        self.title_frame.grid(row=0, column=0, columnspan=3)

        #Create toplevel labels.
        for index,description in enumerate(self.gui_table.columns):
            if (description.lower() != 'type'):
                (Label(self, text=self.item_values[index], foreground='blue')
                       .grid(row=index+1, column=1))
            else:
                (Label(self, text=self.item_values[index], foreground='blue')
                       .grid(row=index+1, column=1))
        self.button_frame.destroy()

        self.button_frame = Frame(self)
        (HoverButton(self.button_frame, text="Confirm", command=self.confirm)
                     .pack(side=LEFT, padx=5, pady=5))
        (HoverButton(self.button_frame, text="Cancel",
                     command=lambda: self.destroy())
                     .pack(side=LEFT, padx=5, pady=5))
        self.button_frame.grid(row=self.grid_size+1, column=0, columnspan=3)


    def confirm(self):
        #Work through Edit_View toplevel to find entry widgets and
        #extract these values to be updated in the given sqlite table.
        #Uses db_update() to update certain column values afterwards and
        #view_products() to display the updated gui table.
        self.changes = []   #list where updated entries will exist
        self.edit_entries = [x for x
                             in reversed(self.grid_slaves())
                             if (x.winfo_class() == 'Entry'
                             or x.winfo_class() == 'TCombobox')]
        for entry in self.edit_entries:
            if entry.winfo_class() == 'Entry' and entry.get():
                #Titlecase before appending.
                self.changes.append(' '.join(word[0].upper() + word[1:]
                                    for word
                                    in entry.get().split()))
            elif entry.winfo_class() == 'TCombobox':
                self.changes.append(entry.get())
            else:
                messagebox.showerror("Input Error",
                                     "At least one input is blank, please try "
                                     + "again.", parent=self)
                return
        self.current_values = [x.cget('text') for x
                               in reversed(self.grid_slaves(column=1))
                               if (x.winfo_class() == 'Label')]
        self.changes = tuple(self.changes + self.current_values)
        db_update()
        edit_db(self.changes, self.sqlite_table, self.gui_table, self.view_fr)
        self.destroy()


class Production_View(Toplevel):
    #Toplevel used to register production.  Subtracts values from raw
    #materials when used.  Handles unfinished products by placing in
    #'in_progress' table to be finished later.
    def __init__(self,master,sqlite_table,gui_table):
        self.master = master
        self.sqlite_table = sqlite_table
        self.gui_table = gui_table
        self.x = (screen_width/2) - (width/2) + 100
        self.y = ((screen_height/2) - (height/2)) + 50
        Toplevel.__init__(self,master=master)

        self.title_frame = Frame(self)
        Label(self.title_frame, text="Production", font="Arial 10 bold").pack()
        self.title_frame.grid(row=0, column=0, columnspan=3, pady=5)

        self.product_frame = Frame(self)
        Label(self.product_frame, text="Total Bottles").grid(row=0, column=0)
        Label(self.product_frame, text="Cases").grid(row=0, column=1)
        Label(self.product_frame, text="Product").grid(row=0, column=2)
        (Entry(self.product_frame, validate='key',
               validatecommand=(self.register(valid_dig),'%S','%d'))
               .grid(row=1, column=0, padx=5))
        (Entry(self.product_frame, validate='key',
               validatecommand=(self.register(valid_dig),'%S','%d'))
               .grid(row=1, column=1, padx=5))
        self.conn = sqlite3.Connection("inventory.db")
        self.conn.row_factory = lambda cursor, row: row[0]
        self.cur = self.conn.cursor()
        self.cur.execute("SELECT product " +
                           "FROM 'bottles'")
        self.product_rows = self.cur.fetchall()
        self.products = ttk.Combobox(self.product_frame,
                                     values=self.product_rows)
        self.products.config(width=20, background="white", justify='center',
                             state='readonly')
        self.products.set(self.product_rows[0])
        self.products.grid(row=1, column=2, padx=5)
        self.product_frame.grid(row=1, column=0, columnspan=3)

        #Raw materials title frame.
        self.materials = Frame(self)
        (Label(self.materials, text="Materials Used", font="Arial 10 bold")
               .pack())
        self.materials.grid(row=3, column=0, columnspan=3, pady=5)

        #Raw materials input frame.
        Label(self, text="Type").grid(row=4, column=0, pady=2)
        Label(self, text="Amount").grid(row=4, column=1, pady=2)
        Label(self, text="Material").grid(row=4, column=2, pady=2)
        self.type_rows = type_options['raw_materials']

        #Create label, entry, option box for each type of raw material.
        for index,description in enumerate(self.type_rows,5):
            Label(self, text=description + ":").grid(row=index, column=0,
                                                     sticky=W)
            (Entry(self, validate='key',
                   validatecommand=(self.register(valid_dig),'%S','%d'))
                   .grid(row=index, column=1))
            self.cur.execute("SELECT product "
                             + "FROM 'raw_materials' " +
                              "WHERE type=\'" + description + "\'")
            self.rows = self.cur.fetchall()
            self.rows.append("None")
            self.opt_menu = ttk.Combobox(self, values=self.rows)
            self.opt_menu.config(width=20, background="white", justify='center',
                                 state='readonly')
            self.opt_menu.set(self.rows[0])
            self.opt_menu.grid(row=index, column=2, padx=5)

        #Finished product checkbox.
        self.grid_size = self.grid_size()[1]
        self.check_var = IntVar()
        self.check_var.set(1)
        self.check_b = Checkbutton(
            self, text="Are the products finished? (i.e. labeled)",
            variable=self.check_var, command=self.cbox_check)
        self.check_b.grid(row=self.grid_size+1, column=0, columnspan=3)

        #Samples input frame.
        self.samples_frame = Frame(self)
        Label(self.samples_frame, text="Samples").grid(row=0, column=0)
        self.samples_entry = Entry(
            self.samples_frame, validate='key',
            validatecommand=(self.register(valid_dig),'%S','%d'))
        self.samples_entry.grid(row=0, column=1)
        self.samples_frame.grid(row=self.grid_size+2, column=0, columnspan=3)

        self.button_frame = Frame(self)
        (HoverButton(self.button_frame, text="Confirm", width=10,
                     command=self.confirm)
                     .pack(side=LEFT, padx=5, pady=5))
        (HoverButton(self.button_frame, text="Cancel", width=10,
                     command=lambda: self.destroy())
                     .pack(side=LEFT, padx=5, pady=5))
        self.button_frame.grid(row=self.grid_size+3, column=0, columnspan=3)

        self.conn.close()
        self.title("Production")
        self.focus()
        self.geometry("+%d+%d" % (self.x, self.y))
        self.resizable(0,0)

    def confirm(self):
        #Updates raw_materials, production_log and bottles/samples or
        #in_progress sqlite tables with corresponding values from
        #production toplevel.
        self.product_amount = (self.product_frame
                               .grid_slaves(row=1, column=0)[0]
                               .get())
        self.case_amount = (self.product_frame
                            .grid_slaves(row=1, column=1)[0]
                            .get())
        self.product_var = (self.product_frame
                            .grid_slaves(row=1, column=2)[0]
                            .get())
        self.samples_var = self.samples_entry.get()
        #Raw material options.
        self.materials = [x.get() for x
                          in reversed(self.grid_slaves())
                          if (x.winfo_class() == 'TCombobox')]
        #Raw material entries.
        self.entries = [x.get() for x
                        in reversed(self.grid_slaves())
                        if (x.winfo_class() == 'Entry')]
        #Raw material product types.
        self.types = [x.cget("text").rstrip(":") for x
                      in reversed(self.grid_slaves())
                      if (x.winfo_class() == 'Label'
                      and x.cget("text").find(":") != -1)]
        for entry,material in zip(self.entries, self.materials):
            #Check material inputs to ensure non-empty values.
            if (not entry and material != "None"):
                messagebox.showerror(
                    "Materials Input Error",
                    "At least one input within the materials used section is "
                    + "blank, please try again.", parent=self)
                return
        #Check product and case amounts.
        if (not self.product_amount) or (not self.case_amount):
            messagebox.showerror(
                "Product Input Error",
                "One or more of the 'Total Bottles' or 'Cases' entries are "
                + "blank, please try again.", parent=self)
            return
        #Check sample amount and 'finished' checkbox.
        if (not self.samples_entry.get()) and (self.check_var.get() == 1):
            messagebox.showerror(
                "Sample Input Error",
                "The samples entry must be non-empty, please try again.",
                parent=self)
            return
        self.curr_date = date.today()
        #Update 'in_progress' table if products checked as unfinished.
        self.conn = sqlite3.Connection("inventory.db")
        self.cur = self.conn.cursor()
        if self.check_var.get() == 0:
            self.cur.execute("INSERT INTO 'in_progress' " +
                                  "VALUES (?,?,?,?)",
                             (self.curr_date, self.product_var,
                              self.product_amount, self.desc_var))
        #Update 'bottles' and 'samples' tables if products are
        #considered finished
        elif self.check_var.get() == 1:
            self.cur.execute("UPDATE 'bottles' " +
                                "SET amount=(amount + ?) " +
                              "WHERE product=?",
                             (self.case_amount, self.product_var))
            self.cur.execute("UPDATE 'samples' " +
                                "SET amount=(amount + ?) " +
                              "WHERE product=?",
                             (self.samples_var, self.product_var))
        #Update 'production log'.
        self.cur.execute("INSERT INTO 'production' " +
                              "VALUES (?,?,?)",
                         (self.curr_date, self.product_var,
                          self.product_amount))
        #Update 'raw_materials' data.
        for (material, subtr,type) in zip(self.materials, self.entries,
                                         self.types):
            if material != "None":
                self.cur.execute("UPDATE 'raw_materials' " +
                                    "SET amount=MAX((amount - ?),0) " +
                                  "WHERE product=? " +
                                    "AND type=?",
                                 (subtr, material,type))
        self.conn.commit()
        self.conn.close()
        db_update()
        view_products('raw_materials', 'All', 'All', raw_tbl)
        self.destroy()

    def cbox_check(self):
        #Activate sample entry box and cases amount entry if checkbox is
        #checked.
        if self.check_var.get() == 1:
            self.samples_entry.config(state='normal')
            self.samples_entry.delete(0, END)
            (self.product_frame.grid_slaves(row=1, column=1)[0]
             .config(state='normal'))
            self.product_frame.grid_slaves(row=1, column=1)[0].delete(0, END)
        #Disable sample entry box and cases amount entry if checkbox is
        #unchecked.
        if self.check_var.get() == 0:
            self.samples_entry.delete(0, END)
            self.samples_entry.insert(0, "0")
            self.samples_entry.config(state='readonly')
            self.product_frame.grid_slaves(row=1, column=1)[0].insert(0, "0")
            (self.product_frame.grid_slaves(row=1, column=1)[0]
             .config(state='readonly'))

            def desc_cancel():
                #Sets cbox value to 1 to prevent user from continuing
                #without entering description value or entering sample
                #amount.
                self.desc_tl.destroy()
                self.check_var.set(1)
                self.cbox_check()


            def desc_set():
                #Set description variable based on description text.
                if  not self.desc_text.compare("end-1c", "==", "1.0"):
                    self.desc_var = self.desc_text.get("1.0", END)
                    self.desc_tl.destroy()
                else:
                    messagebox.showerror("Input Error",
                                         "Please input a description.",
                                         parent=self.desc_tl)
                    return

            #Toplevel to insert description for in_progress table.
            self.desc_var = StringVar()
            self.desc_tl = Toplevel(self)
            (Message(self.desc_tl, text="Please provide a description of why "
                     + "the production was considered unfinished. "
                     + "(ex. 'bottles unlabeled', 'waiting for labels')",
                     width=300)
                     .grid(row=0, column=0, columnspan=2))
            self.desc_text = Text(self.desc_tl, height=2, width=30)
            self.desc_text.grid(row=1, column=0, columnspan=2)
            self.desc_fr = Frame(self.desc_tl)
            self.conf_b = HoverButton(self.desc_fr, text="Confirm",
                                      command=desc_set)
            self.conf_b.grid(row=0, column=0)
            (HoverButton(self.desc_fr, text="Cancel", command=desc_cancel)
                         .grid(row=0, column=1))
            self.desc_fr.grid(row=2, column=0, columnspan=2)
            #Prevent use of 'x-out' button.
            self.desc_tl.protocol("WM_DELETE_WINDOW", disable_event)
            self.desc_tl.title("Production Description")
            self.desc_tl.resizable(0,0)
            self.desc_tl.geometry("+%d+%d" % (self.x + 30, self.y + 30))
            self.desc_tl.focus()
            #Prevent user from clicking outside of toplevel.
            self.desc_tl.grab_set()


class Purchase_Order_View(Toplevel):
    #Toplevel used to create purchase orders.  Updates purchase_order
    #and bottles/samples based on values retrieved from the toplevel.
    def __init__(self, master):
        self.master = master
        self.x = (screen_width/2) - 350
        self.y = (screen_height/2) - 350
        self.conn = sqlite3.Connection("inventory.db")
        self.conn.row_factory = lambda cursor, row: row[0]
        self.cur = self.conn.cursor()
        self.cur.execute("SELECT product " +
                           "FROM 'bottles'")
        self.product_rows = self.cur.fetchall()
        self.conn.close()
        Toplevel.__init__(self, master=self.master)

        self.title_fr = Frame(self)
        Label(self.title_fr, text="Purchase Order", font="Arial 10 bold").pack()
        self.title_fr.grid(row=0, column=0, columnspan=2)

        #Frame containing purchase order shipment information.
        self.info_fr = Frame(self, pady=2)
        Label(self.info_fr, text="From:").grid(row=0, column=0, sticky=W)
        Label(self.info_fr, text="PO Number:").grid(row=1, column=0, sticky=W)
        Label(self.info_fr, text="To:").grid(row=2, column=0, sticky=W)
        Label(self.info_fr, text="PO Date:").grid(row=0, column=2, sticky=W)
        Label(self.info_fr, text="Pick Up Date:").grid(row=1, column=2,
                                                       sticky=W)
        Entry(self.info_fr, justify='center').grid(row=0, column=1)

        #Search for last purchase order in corresponding year folder
        #if none exists, create this year's folder.
        self.year = datetime.now().year
        try:
            self.files = (os.listdir(os.getcwd() + "\\purchase_orders\\"
                          + str(self.year)))
        except FileNotFoundError:
            os.mkdir(os.getcwd() + "\\purchase_orders\\" + str(self.year))
            self.files = []
        self.po_nums = []
        #Update po-number entry with latest po number + 1.
        if self.files:
            for file in self.files:
                self.mo = poRegex.search(file)
                if self.mo:
                    self.po_nums.append(int(self.mo.group(3)))
            self.new_po_num = (str(self.year) + "-" # ex 2019-005
                               + '{:03}'.format(max(self.po_nums) + 1))
        else:
            # Make YYYY-001
            self.new_po_num = str(self.year) + "-" + '{:03}'.format(1)
        self.po_entry = Entry(self.info_fr, justify='center')
        self.po_entry.insert(0, self.new_po_num)
        self.po_entry.grid(row=1, column=1)
        Entry(self.info_fr, justify='center').grid(row=2, column=1)
        self.po_date = Entry(self.info_fr, justify='center', state='readonly')
        self.po_date.grid(row=0, column=3)
        self.po_cal_link = HoverButton(self.info_fr, image=cal_photo,
                                       command=lambda:
                                       cal_button(self, self.po_date))
        self.po_cal_link.image = cal_photo
        self.po_cal_link.grid(row=0,column=4)
        self.pu_date = Entry(self.info_fr, justify='center', state='readonly')
        self.pu_date.grid(row=1, column=3)
        self.pu_cal_link = HoverButton(self.info_fr, image=cal_photo,
                                       command=lambda:
                                       cal_button(self, self.pu_date))
        self.pu_cal_link.image = cal_photo
        self.pu_cal_link.grid(row=1,column=4)
        self.info_fr.grid(row=1, column=0, columnspan=2)

        #Frame containing purchase order product information.
        self.order_fr = Frame(self, padx=33)
        Label(self.order_fr, text="QTY").grid(row=0, column=0, sticky=N+E+S+W)
        Label(self.order_fr, text="UNIT").grid(row=0, column=1, sticky=N+E+S+W)
        Label(self.order_fr, text="PRODUCT").grid(row=0, column=2,
                                                  sticky=N+E+S+W)
        Label(self.order_fr, text="UNIT COST").grid(row=0, column=3,
                                                    sticky=N+E+S+W)
        Label(self.order_fr, text="TOTAL").grid(row=0, column=4, sticky=N+E+S+W)
        for i in range(1,19):   #Purchase order product information.
            (Entry(
                self.order_fr, width=5, justify="center", validate="key",
                validatecommand=(self.register(valid_dig),'%S','%d'))
                .grid(row=i, column=0, sticky=N+E+S+W))
            (ttk.Combobox(
                self.order_fr, values=['Cases','Bottles'],
                width=7, justify="center", state='readonly')
                .grid(row=i, column=1, sticky=N+E+S+W))
            (ttk.Combobox(
                self.order_fr, values=self.product_rows, justify="center",
                state='readonly')
                .grid(row=i, column=2))
            (Entry(
                self.order_fr, width=12, justify="center", validate="key",
                validatecommand=(self.register(valid_dec),'%S','%s','%d'))
                .grid(row=i, column=3, sticky=N+E+S+W))
            (Entry(
                self.order_fr, width=12, justify="center", bg="light gray")
                .grid(row=i,column=4,sticky=N+E+S+W))
        (Label(
            self.order_fr, text="TOTAL", background="dark slate gray",
            relief="raised", fg="white")
            .grid(row=19, column=0, columnspan=5, sticky=E+W))
        self.total_var = StringVar()
        self.total_label = Label(self.order_fr, background="gray", fg="white",
                                 width=10, relief="raised",
                                 textvariable=self.total_var)
        self.total_label.grid(row=19, column=4, sticky=E+W)
        for label in self.order_fr.grid_slaves(row=0):
            label.config(background="dark slate gray", relief="raised",
                         fg="white")
        self.order_fr.grid(row=2, column=0, columnspan=2, pady=2)

        self.check_var = IntVar()
        self.check_var.set(1)
        self.check_b = Checkbutton(self, text="Pending Purchase Order",
                                   variable=self.check_var)
        self.check_b.grid(row=3, column=0, columnspan=2)

        self.btn_fr = Frame(self)
        (HoverButton(self.btn_fr, text="Confirm", command=self.confirm)
                     .grid(row=0, column=0, padx=10))
        (HoverButton(self.btn_fr, text="Cancel", command=lambda: self.destroy())
                     .grid(row=0, column=1, padx=10))
        self.btn_fr.grid(row=4, column=0, columnspan=2, pady=2)

        self.total_after()
        self.geometry("%dx%d+%d+%d" % (464, 625, self.x, self.y))
        self.focus()

    def total_after(self):
        #Updates total column entry values to be product of quantity and
        #price.  Sums total columns into final total column, total_var
        self.total_entries = [x for x
                              in reversed(self.order_fr.grid_slaves(column=4))
                              if x.winfo_class() == 'Entry']
        self.total_sum = 0
        for entry,i in zip(self.total_entries,range(1, 19)):
            entry.delete(0, END)
            try:
                # (quantity * unit cost)
                self.row_total = (
                    float(self.order_fr.grid_slaves(row=i, column=0)[0].get())
                    *float(self.order_fr.grid_slaves(row=i, column=3)[0].get()))
                self.row_amt = "{0:,.2f}".format(self.row_total)
                entry.insert(0,"$%s" % self.row_amt)
                self.total_sum += self.row_total
            except:
                pass
        self.total_sum = "{0:,.2f}".format(self.total_sum)
        self.total_var.set("$%s" % self.total_sum)
        self.after_func = self.after(150, self.total_after)

    def confirm(self):
        self.open_ques = messagebox.askquestion(
            "Purchase Order Confirmation",
            "Are you sure you want to confirm? Please make sure everything is "
            + "entered correctly. Confirming will update inventory values and "
            + "save the purchase order with the file name, "
            + self.new_po_num + ".xlsx. \n \n"
            + "If the 'Pending Purchase Order' checkbox is checked, "
            + "the purchase order will be stored within the 'Pending' tab to "
            + "be completed or removed at a later date. \n \n"
            + "***IMPORTANT***  The PO-Number is calculated from completed "
            + "purchase orders. If there are any pending purchase orders, "
            + "please fulfill or cancel them before continuing. You may also "
            + "update the purchase order number yourself.", parent=self)

        if self.open_ques == 'no':
            return self.total_after()
        else:
            #Stop after_func to ensure total_var value.
            self.after_cancel(self.after_func)
            self.info_entries = [x.get() for x
                                 in reversed(self.info_fr.grid_slaves())
                                 if x.winfo_class() == 'Entry']
            for entry in self.info_entries:
                if not entry:
                    messagebox.showerror(
                        "Input Error",
                        "Please make sure all of the top entries have values.",
                        parent=self)
                    return self.total_after()
            self.po_entries = [x.get() for x
                               in reversed(self.order_fr.grid_slaves())
                               if (x.winfo_class() == 'Entry'
                               or x.winfo_class() == "TCombobox")]
            #List of lists containing po order values.
            self.complete_po_lists = [self.po_entries[x:x+5] for x
                                      in range(0,len(self.po_entries),5)]
            for list in self.complete_po_lists:
                if any(list) and not all(list):
                    messagebox.showerror(
                        "Input Error",
                        "Please make sure all of the purchase order entries "
                        + "are fully complete.", parent=self)
                    return self.total_after()
                else:
                    continue

            if self.check_var.get() == 1:
                #Add purchase orders to 'pending_po' table.
                self.conn = sqlite3.Connection("inventory.db")
                self.cur = self.conn.cursor()
                for po_list in (x for x in self.complete_po_lists if all(x)):
                    self.cur.execute("INSERT INTO 'pending_po' " +
                                          "VALUES (?,?,?,?,?,?,?,?,?)",
                                     (self.info_entries[3],
                                      self.info_entries[4], po_list[2],
                                      po_list[0], po_list[1], po_list[3],
                                      po_list[4], self.info_entries[2],
                                      self.po_entry.get()))
                self.conn.commit()
                self.conn.close()
            else:
                confirm_po(po_vfr, self.info_entries,
                           self.complete_po_lists, self.po_entry.get())
            self.destroy()


class Finish_View(Toplevel):
    #Toplevel used to finish 'in progress' production values.  Deletes
    #in_progress table entry and updates bottles/samples upon confirmation.
    def __init__(self, master, values):
        self.master = master
        self.values = values
        self.x = x + 150
        self.y = y + 20
        self.product = self.values[1]
        Toplevel.__init__(self, master=self.master)

        self.title_fr = Frame(self)
        (Label(self.title_fr, text="Finish Production", font="Arial 10 bold",
               pady=5)
               .pack())
        self.title_fr.grid(row=0, column=0, columnspan=2, pady=5)

        self.prod_fr = Frame(self)
        (Label(self.prod_fr, text=self.product)
               .grid(row=0, column=0, columnspan=2))
        (Label(self.prod_fr, text="Cases:")
               .grid(row=1, column=0, sticky=W, pady=2))
        (Entry(self.prod_fr, validate="key",
               validatecommand=(self.register(valid_dig),'%S','%d'))
               .grid(row=1, column=1, pady=2))
        (Label(self.prod_fr, text="Samples:")
               .grid(row=2, column=0, sticky=W, pady=2))
        (Entry(self.prod_fr, validate="key",
               validatecommand=(self.register(valid_dig),'%S','%d'))
               .grid(row=2, column=1, pady=2))
        self.prod_fr.grid(row=1, column=0, columnspan=2)

        self.button_fr = Frame(self)
        (HoverButton(self.button_fr, text="Confirm", command=self.confirm)
                     .pack(side=LEFT))
        (HoverButton(self.button_fr, text="Cancel",
                     command=lambda: self.destroy())
                     .pack(side=LEFT))
        self.button_fr.grid(row=2, column=0, columnspan=2)

        self.title("Finish " + self.product + " Production")
        self.geometry("%dx%d+%d+%d" % (178, 140, self.x, self.y))
        self.resizable(0,0)
        self.focus()

    def confirm(self):
        #Update 'bottles' and 'samples' tables with newly created
        #products. Remove info from 'in_progress' table.
        self.conf_quest = messagebox.askquestion(
            "Finish this product?",
            "Are you sure you want to confirm? Make sure everything is entered "
            + "correctly before continuing.", parent=self)
        if self.conf_quest == "yes":
            self.entries = [x.get() for x
                            in reversed(self.prod_fr.grid_slaves())
                            if x.winfo_class() == "Entry"]
            if all(self.entries):
                self.conn = sqlite3.Connection("inventory.db")
                self.cur = self.conn.cursor()
                self.cur.execute("UPDATE 'bottles' " +
                                    "SET amount=(amount + ?) " +
                                  "WHERE product=?",
                                 (self.entries[0], self.product))
                self.cur.execute("UPDATE 'samples' " +
                                    "SET amount=(amount + ?) " +
                                  "WHERE product=?",
                                  (self.entries[1], self.product))
                self.cur.execute("DELETE FROM 'in_progress' " +
                                       "WHERE product=?",
                                 (self.product,))
                self.conn.commit()
                self.conn.close()
                db_update()
                view_products('in_progress', 'All', 'All', inprog_tbl)
                self.destroy()
            else:
                messagebox.showerror(
                    "Input Error",
                    "Please make sure all of the info entries have values.",
                    parent=self)
                return
        else:
            return


class Mash_Production_View(Toplevel):
    #Toplevel to input values for mash production.
    def __init__(self, master, mash_table):
        self.master = master
        self.mash_table = mash_table
        self.x = x + 150
        self.y = y + 150
        #Populated with lists of length 3 (grain, amt, order #) in grain_recur
        self.grain_info_tbl = []
        Toplevel.__init__(self, master=self.master)
        #Get previous mash information.
        try:
            self.conn = sqlite3.Connection("inventory.db")
            self.cur = self.conn.cursor()
            self.cur.execute("SELECT mash_no, type " +
                               "FROM mashes " +
                           "ORDER BY mash_no " +
                         "DESC LIMIT 1")
            self.prev_mash = list(self.cur)[0]
            self.prev_mash_num = self.prev_mash[0]
            self.prev_mash_type = self.prev_mash[1]
            self.conn.close()
            #Mash number regex matches.
            self.mo = mashRegex.search(self.prev_mash_num)
            self.year = self.mo.group(1)    #Prev mash's year.
            self.mash_count = self.mo.group(5)  #Prev mash's ID number.
            self.mash_letter = self.mo.group(6) #Prev mash's letter variable.
            self.mash_letters = list(string.ascii_uppercase[:8]) #Letters A-H
        except:
            self.year = int(datetime.now().year)
            self.mash_count = 0
            self.mash_letter = "A"
            self.prev_mash_type = None
        self.conn = sqlite3.Connection("inventory.db")
        self.cur = self.conn.cursor()
        self.cur.execute("SELECT type, order_number " +
                           "FROM grain")
        self.grain_ord_nums = self.cur.fetchall()
        self.conn.close()
        self.grain_ord_dict = {}
        for (key, value) in self.grain_ord_nums:
            if key in self.grain_ord_dict:
                self.grain_ord_dict[key].append(value)
            else:
                self.grain_ord_dict[key] = [value]

        self.title_fr = Frame(self)
        (Label(self.title_fr, text="Mash Production", font="Arial 10 bold",
               pady=5)
               .pack())
        self.title_fr.grid(row=0, column=0, columnspan=3)

        self.type_fr = Frame(self)
        Label(self.type_fr, text="Mash Type:").grid(row=0, column=0)
        self.type_menu = ttk.Combobox(
            self.type_fr, values=["Bourbon","Rye","Malt","Rum"],
            width=16, justify="center", state="readonly")
        self.type_menu.set("Bourbon")
        self.type_menu.bind("<<ComboboxSelected>>", self.tplvl_upd)
        self.type_menu.grid(row=0, column=1)
        Label(self.type_fr, text="Mash Number:").grid(row=1, column=0)
        self.mash_num_entry = Entry(self.type_fr, justify='center')
        self.mash_num_entry.grid(row=1, column=1)
        Label(self.type_fr, text="Date:").grid(row=2, column=0)
        self.date = StringVar()
        #Update mash number with month value.
        self.date.trace(
            "w",
            lambda name, index, mode:
            self.mash_num_upd(self.prev_mash_type, self.type_menu.get()))
        self.date_entry = Entry(self.type_fr, state="readonly",
                                justify="center", textvariable=self.date)
        self.date_entry.grid(row=2,column=1)
        self.cal_link = HoverButton(self.type_fr, image=cal_photo,
                                    command=lambda:
                                    cal_button(self, self.date_entry))
        self.cal_link.image = cal_photo
        self.cal_link.grid(row=2, column=2)
        self.type_fr.grid(row=1, column=0, columnspan=3)

        self.grain_fr = Frame(self, pady=5, padx=5, height=100, width=340)
        self.grain_fr.grid_propagate(0) #Frame size doesn't change

        self.button_fr = Frame(self, padx=10)
        (HoverButton(self.button_fr, text="Confirm", command=self.confirm)
                     .grid(row=0, column=0))
        (HoverButton(self.button_fr, text="Cancel",
                     command=lambda: self.destroy())
                     .grid(row=0, column=1))
        self.button_fr.grid(row=3, column=0, columnspan=3)

        self.title("Mash Production")
        self.geometry("%dx%d+%d+%d" % (350, 240, self.x, self.y))
        self.resizable(0,0)
        self.focus()
        self.type_menu.event_generate("<<ComboboxSelected>>")

    def mash_num_upd(self, prev_type, curr_type):
        #Update mash number entry based on current grain type and
        #previous mash type.
        self.mash_num_entry.delete(0, END)
        if self.date_entry.get():
            self.month = self.date_entry.get()[5:7]
        else:
            self.month = '{:02d}'.format(datetime.now().month)
        #ex 2019/03-4A
        self.new_batch_num = (str(self.year) + "/" + str(self.month) + "-"
                              + str(int(self.mash_count) + 1) + "A")
        #Handle new year case.
        if int(self.year) != int(datetime.now().year):
            self.mash_num_entry.insert(0, self.year + "/1-1A")
        else:
            if prev_type == curr_type:
                if self.mash_letter != "H": #Same type, same batch case.
                    self.mash_let_indx = (self.mash_letters
                                          .index(self.mash_letter) + 1)
                    self.new_batch_num = (
                        str(self.year) + "/" + str(self.month)
                        + "-" + str(self.mash_count)
                        + self.mash_letters[self.mash_let_indx])
                    self.mash_num_entry.insert(0, self.new_batch_num)
                else:   #Same type, next batch case.
                    self.mash_num_entry.insert(0, self.new_batch_num)
            else:   #New type, new batch case.
                self.mash_num_entry.insert(0, self.new_batch_num)

    def fill_frame(self, gr_lst):
        #Update grain frame with grain types used for selected mash
        #type.
        for index, grain in enumerate(gr_lst,1):
            Label(self.grain_fr, text=grain).grid(row=index, column=0)
            (Entry(self.grain_fr, validate='key',
                   validatecommand=(self.register(valid_dig),
                                    '%S', '%d'))
                   .grid(row=index, column=1))
            (ttk.Combobox(self.grain_fr,
                          values=self.grain_ord_dict[grain], width=16,
                          justify='center', state='readonly')
                          .grid(row=index, column=2, padx=3))

    def tplvl_upd(self, event):
        #Remove grain inputs and replace with new ones corresponding to
        #grain type.  Update mash number based on previous and current
        #mash types
        self.type = self.type_menu.get()
        self.mash_num_upd(self.prev_mash_type, self.type)
        for widg in self.grain_fr.grid_slaves():
            widg.grid_forget()
        (Label(self.grain_fr, text="Grain", font="Arial 10 bold")
               .grid(row=0, column=0, columnspan=3))
        try:
            if self.type == "Bourbon":
                self.fill_frame(["Corn","Rye","Malted Barley"])
            elif self.type == "Rye":
                self.fill_frame(["Rye","Malted Wheat"])
            elif self.type == "Malt":
                self.fill_frame(["Malted Barley","Wheat","Oat"])
            else:
                self.fill_frame(["Molasses"])
        except KeyError:
            self.destroy()
            messagebox.showerror(
                "Grain Inventory Error",
                "There was an issue retrieving the purchase order number for " +
                "one of the grain types required for this type of mash. " +
                "Please make sure there is an inventory value for each type " +
                "of grain needed in this mash. \n \n" +
                "If you are receiving this message upon opening " +
                "'Produce Mash', there is a missing grain for Bourbon.")
            return
        self.grain_fr.grid(row=2, column=0, columnspan=3)

    def confirm(self):
        #Subtract grain used from 'grain' table, update 'mash_log' based
        #on entries. Create a 'mash log' word file with certain info
        #filled out.
        self.grain_types = [x.cget("text") for x
                            in reversed(self.grain_fr.grid_slaves())
                            if x.winfo_class() == "Label"][1:]
        self.grain_amts = [x.get() for x
                              in reversed(self.grain_fr.grid_slaves())
                              if x.winfo_class() == "Entry"]
        self.type_entries = [x.get() for x
                             in reversed(self.type_fr.grid_slaves())
                             if x.winfo_class() == "Entry"]
        self.order_nums = [x.get() for x
                           in reversed(self.grain_fr.grid_slaves())
                           if x.winfo_class() == "TCombobox"]
        self.entry_check = self.grain_amts + self.type_entries + self.order_nums
        for entry in self.entry_check:
            if not entry:
                messagebox.showerror(
                    "Input Error",
                    "Please make sure all of the entries are fully " +
                    "complete.", parent=self)
                return
        try:
            self.file_path = os.getcwd()
            self.file = open(self.file_path +
                             '/production_sheets/Mash_Log.docx', 'rb')
            self.document = Document(self.file)
            self.file.close()
        except:
            messagebox.showerror(
                "File Error",
                "It seems the Word Document you are trying to edit or change "
                + "is already open, please close it and try again.",
                parent=self)
            return

        #Subtract grain amounts from inventory.
        self.conn = sqlite3.Connection("inventory.db")
        self.cur = self.conn.cursor()
        for type, amount, order_num in zip(self.grain_types, self.grain_amts,
                                self.order_nums):
            self.grain_recur(type, amount, order_num)
        self.cur.execute("INSERT INTO mashes " +
                              "VALUES (?,?,?,?)",
                         (self.date_entry.get(), self.mash_num_entry.get(),
                         self.type_menu.get(),
                         ", ".join(self.order_nums)))
        self.conn.commit()
        self.conn.close()

        #Word doc tables
        self.info_table = self.document.tables[0]
        self.grain_table = self.document.tables[1]
        #Write info for top of file
        for (row, info) in zip(self.info_table.rows,
                              [self.date_entry.get(),
                               self.type_menu.get(),
                               self.mash_num_entry.get()]):
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = info
                    for run in para.runs:
                        run.font.name = "Verdana"
                        run.font.size = Pt(14)
        #Write info for grain table section of file
        for (row, gr_list) in zip(self.grain_table.rows, self.grain_info_tbl):
            for (cell, num) in zip(row.cells, range(3)):
                for para in cell.paragraphs:
                    para.text = gr_list[num]
                    for run in para.runs:
                        run.font.name = "Verdana"
                        run.font.size = Pt(14)

        self.document.save(self.file_path
                           + "/last_mash/Last_Mash_Log.docx")

        self.open_ques = messagebox.askquestion(
            "Open the Mash Word File?",
            "Would you like to open the Mash Log file in Word? This will "
            + "allow you to print it now.")
        if self.open_ques == "yes":
            try:
                os.system('start ' + self.file_path
                          + "/last_mash/Last_Mash_Log.docx")
            except:
                messagebox.showerror(
                    "Program Error",
                    "There was an error opening Word.", parent=self)
        else:
            pass
        db_update()
        view_products('grain', 'All', 'All', grain_tbl)
        self.destroy()

    def grain_recur(self, type, amount, order_num, first=True):
        #Subtract amounts from respective grain. Recursion occurs when
        #a grain amount is zeroed out. User is asked for a second order
        #number for that type of grain to subtract from.
        self.cur.execute("SELECT amount, date " +
                           "FROM grain " +
                          "WHERE type=? AND order_number=?",
                         (type, order_num))
        self.grain_lst = list(self.cur)
        print(self.grain_lst)
        self.grain_amt = (self.grain_lst[0][0])
        self.grain_date = (self.grain_lst[0][1])
        if self.grain_amt:
            self.grain_diff = int(self.grain_amt) - int(amount)
            if self.grain_diff > 0:
                self.cur.execute("UPDATE grain " +
                                    "SET amount=? " +
                                  "WHERE type=? " +
                                    "AND order_number=?",
                                 (self.grain_diff, type, order_num))
            else:
                self.cur.execute("DELETE FROM grain " +
                                       "WHERE type=? " +
                                         "AND order_number=?",
                                 (type, order_num))
                self.cur.execute("INSERT INTO grain_log " +
                                      "VALUES (?,?,?,?)",
                                 (self.grain_date, self.date_entry.get(),
                                  type, order_num))
            if first == True:
                self.grain_info_tbl.append([type,
                                            str(amount),
                                            str(order_num)])
            elif first == False:
                for lst in self.grain_info_tbl:
                    print(lst)
                    if lst[0] == type:
                        print(lst[0])
                        lst[2] = lst[2] + ", " + order_num
                        self.order_nums.append(order_num)
            if self.grain_diff < 0:
                self.grain_diff = abs(self.grain_diff)
                self.grain_recur_tplvl(type, order_num)
        else:
            messagebox.showerror(
                "Grain Error",
                "There doesn't seem to be an inventory value for " + type +
                ", or there isn't enough grain, please fix this and try again.",
                parent=self)
            self.conn.close()
            raise ValueError("Grain Error")

    def grain_recur_tplvl(self, type, order_num):
        #Toplevel to add another order number to fulfill grain amount
        self.combox_values = [x for x
                              in self.grain_ord_dict[type]
                              if x != order_num]
        if self.combox_values:
            self.recur_tplvl = Toplevel(self)
            (Label(self.recur_tplvl, text = "There is " + str(self.grain_diff)
                   + "lbs. " + "of " + type + " left to be used, \n " +
                  "please select another order number to subtract it from.")
                  .grid(row=0, column=0, columnspan=2, pady=5))
            self.new_ord_box = ttk.Combobox(self.recur_tplvl,
                                            values=self.combox_values)
            self.new_ord_box.config(width=16, background="white",
                                      justify='center', state='readonly')
            self.new_ord_box.set(self.combox_values[0])
            self.new_ord_box.grid(row=1, column=0, columnspan=2, pady=5)
            (HoverButton(self.recur_tplvl, text="Confirm",
                         command=lambda: self.recur_confirm(type))
                         .grid(row=2, column=0, pady=5))
            (HoverButton(self.recur_tplvl, text="Cancel",
                         command=lambda: self.recur_cancel())
                         .grid(row=2, column=1, pady=5))

            self.recur_tplvl.protocol("WM_DELETE_WINDOW", disable_event)
            self.recur_tplvl.title("Production Description")
            self.recur_tplvl.resizable(0,0)
            self.recur_tplvl.geometry("+%d+%d" % (self.x + 60, self.y + 60))
            self.recur_tplvl.focus()
            self.recur_tplvl.grab_set()
            self.wait_window(self.recur_tplvl)
        else:
            messagebox.showerror(
                "Grain Error",
                "There doesn't seem to be an inventory value for " + type +
                ", or there isn't enough grain, please fix this and try again.",
                parent=self)
            self.conn.close()
            raise ValueError("Grain Error")

    def recur_cancel(self):
        #Button function to destroy toplevel and close sqlite connection
        self.recur_tplvl.destroy()
        self.conn.close()
        raise ValueError("Grain Error")

    def recur_confirm(self, type):
        #Button function to destroy toplevel and run grain_recur with
        #new inputs
        self.grain_recur(type, self.grain_diff, self.new_ord_box.get(),
                         first=False)
        self.recur_tplvl.destroy()

class Reports_Frame(Frame):
    #Creates frame containing inventory information from each sqlite
    #table.
    def __init__(self, master):
        self.master = master
        self.cur_year = datetime.now().year
        self.cur_month_ind = datetime.now().month
        self.year_choices = list(range(2019, self.cur_year + 1))
        self.month_choices = list(calendar.month_abbr)
        Frame.__init__(self, master)
        monthly_reports_update()

        self.year_fr = LabelFrame(self, text="YEAR", font="Arial 8 bold")
        self.year_cmbo_box = ttk.Combobox(
            self.year_fr, values=self.year_choices, state='readonly',
            justify='center')
        self.year_cmbo_box.set(self.cur_year)
        self.year_cmbo_box.bind("<<ComboboxSelected>>", self.year_upd)
        self.year_cmbo_box.pack(padx=5, pady=5)
        self.year_fr.grid(row=0, column=0, padx=5, pady=5)

        self.month_fr = LabelFrame(self, text="MONTH", font="Arial 8 bold")
        self.month_cmbo_box = ttk.Combobox(
            self.month_fr, values=self.month_choices, state='readonly',
            justify='center')
        self.month_cmbo_box.set(self.month_choices[self.cur_month_ind])
        self.month_cmbo_box.bind("<<ComboboxSelected>>", self.month_upd)
        self.month_cmbo_box.pack(padx=5, pady=5)
        self.month_fr.grid(row=0, column=1, padx=5, pady=5)

        self.invent_fr = LabelFrame(self, text="Inventory",
                                   font="Arial 12 bold", fg="dark slate gray")
        self.barrel_fr = LabelFrame(self, text="Barrels",
                                    font="Arial 12 bold", fg="dark slate gray")
        self.po_fr = LabelFrame(self, text="Purchase Orders",
                                font="Arial 12 bold", fg="dark slate gray")

        self.logo_fr = LabelFrame(self, font="Arial 10 bold",
                                  fg="dark slate gray")
        self.logo_path = "ADCO_Logo.jpg"
        self.img = Image.open(self.logo_path)
        self.img = self.img.resize((640,520))
        self.img = ImageTk.PhotoImage(self.img)
        Label(self.logo_fr, image=self.img).grid(row=0, column=0)
        self.logo_fr.grid(row=0, column=2, rowspan=5, sticky="nesw", pady=3,
                          padx=10)

        self.year_cmbo_box.event_generate("<<ComboboxSelected>>")

    def year_upd(self, event):
        #Generate combobox selection event for months, which will make
        #changes to values displayed.
        monthly_reports_update()
        self.year_sel = self.year_cmbo_box.get()
        self.month_cmbo_box.event_generate("<<ComboboxSelected>>")

    def month_upd(self, event):
        #Retrieve and display values from 'monthly_reports' table based
        #on year and month selected.
        self.month_sel = self.month_cmbo_box.get()
        self.month_sel = self.month_choices.index(self.month_sel)
        self.month_sel = "{:02}".format(self.month_sel)
        self.date_sel = ("%s-%s") % (self.year_sel, self.month_sel)
        for widg in [*self.invent_fr.grid_slaves(), # * unpacks contents
                     *self.barrel_fr.grid_slaves(), # of iterables
                     *self.po_fr.grid_slaves()]:
            widg.grid_forget()
        #Inventory Values
        self.conn = sqlite3.Connection("inventory.db")
        self.conn.row_factory = lambda cursor, row: row[1:3]
        self.cur = self.conn.cursor()
        self.cur.execute("SELECT * " +
                           "FROM monthly_reports " +
                          "WHERE date=?",
                          (self.date_sel,))
        self.monthly_totals = [list(x) for x in self.cur.fetchall()]
        if self.monthly_totals:
            self.invent_vals_positions = {'raw_materials':0,
                                          'pending_cogs':1,
                                          'bottles':2,
                                          'samples':3,
                                          'pending_po':4,
                                          'grain':5}
            self.invent_vals = [x for x
                               in self.monthly_totals
                               if x[0] in self.invent_vals_positions.keys()]
            self.invent_vals.sort(key=lambda x:
                                  self.invent_vals_positions[x[0]])
            self.barrel_vals_positions = {'barreled_rum':0,
                                          'barreled_whiskey':1,
                                          'barrels':2}
            self.barrel_vals = [x for x
                                in self.monthly_totals
                                if x[0] in self.barrel_vals_positions.keys()]
            self.barrel_vals.sort(key=lambda x:
                                  self.barrel_vals_positions[x[0]])
            self.po_vals_positions = {'purchase_order_sales':0,
                                      'purchase_order_cogs':1,
                                      'pending_sales':2,
                                      'pending_cogs':3}
            self.po_vals = [x for x
                            in self.monthly_totals
                            if x[0] in self.po_vals_positions.keys()]
            self.po_vals.sort(key=lambda x: self.po_vals_positions[x[0]])
            self.conn.close()
            self.monthly_frames_fill(self.invent_vals, self.invent_fr)
            self.monthly_frames_fill(self.barrel_vals, self.barrel_fr)
            self.monthly_frames_fill(self.po_vals, self.po_fr)
        else:
            (Label(self.invent_fr, text="N/A", font="Arial 30 bold", width=15,
                   fg="gray")
                   .grid(row=0, column=0, columnspan=2))
            (Label(self.barrel_fr, text="N/A", font="Arial 30 bold",
                   fg="gray")
                   .grid(row=0, column=0, columnspan=2))
            (Label(self.po_fr, text="N/A", font="Arial 30 bold",
                   fg="gray")
                   .grid(row=0, column=0, columnspan=2))

        self.invent_fr.grid(row=1, column=0, columnspan=2, padx=5, pady=5,
                            sticky="NESW")
        self.barrel_fr.grid(row=2, column=0, columnspan=2, padx=5, pady=5,
                            sticky="NESW")
        self.po_fr.grid(row=3, column=0, columnspan=2, padx=5, pady=5,
                        sticky="NESW")

    def monthly_frames_fill(self, inv_vals, inv_fr):
        self.invent_sum = 0
        self.grid_ind = 0
        self.neg_vals = ['purchase_order_cogs', 'pending_cogs']
        for index, lst in enumerate(inv_vals):
            self.invent_sum += float(lst[1])
            self.grid_ind += 1
            #ex. Purchase Orders:
            self.txt1 = str(lst[0].replace("_"," ").upper() + ":")
            self.txt2 = "${0:,.2f}".format(lst[1]).replace("-","")
            (Label(inv_fr, text=self.txt1, font="Arial 10 bold")
                   .grid(row=index, column=0, padx=20, sticky="W"))
            self.total_label = Label(inv_fr, text=self.txt2,
                                     font="Arial 10 bold", borderwidth=1,
                                     relief="solid", width=9)
            if lst[1] < 0:
                self.total_label.config(bg='pink', fg='red')
            self.total_label.grid(row=index, column=1, ipadx=20, sticky="E")
        (Label(inv_fr, text="TOTAL:", font="Arial 12 bold")
               .grid(row=self.grid_ind, column=0, ipadx=20, sticky="W"))
        (Label(inv_fr, text="${0:,.2f}".format(self.invent_sum),
               font="Arial 12 bold", borderwidth=2, relief="solid",
               width=10)
               .grid(row=self.grid_ind, column=1, ipadx=20, sticky="E"))
        # Force widget to fill column
        inv_fr.columnconfigure(0, weight=1)
        inv_fr.columnconfigure(1, weight=1)


class Sheet_Label(Label):
    #Creates a clickable label with link to file in given file location.
    def __init__(self, master, text, file_location):

        Label.__init__(self, master, text=text, cursor="hand2",
                       font="Times 14 underline", fg="#0000EE")
        def button_click(event):
            #Changes label color from 'blue' to 'purple' and opens the
            #file.
            if self['fg'] =="#0000EE":
                self['fg'] = "#551A8B"
            else:
                self['fg'] = "#551A8B"
            file = webbrowser.open_new(file_location)
        self.bind("<Button-1>", func=button_click)


class Command_Frame(Frame):
    #Creates frame on the left side of the treeview tables.  Used to
    #place command buttons for interacting with data.
    def __init__(self,master):
        self.master = master
        self.height = height
        self.width = command_width
        Frame.__init__(self, master=self.master, height=self.height,
                       width=self.width)


class View_Frame(LabelFrame):
    #Frame that contains comboboxes for selecting filters within the
    #current displayed inventory table. Filters will cause the table
    #to display updated values based on the chosen filter.
    def __init__(self, master, sqlite_table, gui_table):
        self.master = master
        self.sqlite_table = sqlite_table
        self.gui_table = gui_table
        self.height = int(screen_height/1.5)
        self.text_var = StringVar()
        LabelFrame.__init__(self, master=self.master, height=self.height, bd=5,
                            relief=RIDGE, text="View", font="bold")

        #Columns selector.
        Label(self, text="Columns").grid(row=0, column=0)
        self.column_vals = [x for x
                            in list(self.gui_table["columns"])
                            if (x not in [
                                "Price", "Total", "Amount", "Unit", "Case Size",
                                "Grains", "Mash No", "PO No.", "Proof Gallons",
                                "Order No"])
                            ]
        self.column_vals.append("All")
        self.columns = ttk.Combobox(self, values=self.column_vals)
        self.columns.set(self.column_vals[-1])
        self.columns.config(width=18, background="white", justify='center',
                            state='readonly')
        self.columns.bind("<<ComboboxSelected>>", self.col_upd)
        self.columns.grid(row=1, column=0, pady=3, padx=5)

        #Item selector.
        Label(self, text="Values").grid(row=2, column=0)
        self.rows = ttk.Combobox(self, values=["N/A"])
        self.rows.set("N/A")
        self.rows.config(width=18, background="white", justify='center',
                         state='readonly')
        self.rows.bind("<<ComboboxSelected>>", self.row_upd)
        self.rows.grid(row=3, column=0, pady=3, padx=5)

        #Total value label.
        if any(s in self.gui_table.columns for s in ('Price', 'Proof Gallons')):
            if self.sqlite_table in ["purchase_orders","pending_po"]:
                self.inv_label = "Transactions Value"
            else:
                self.inv_label = "Inventory Value"
            Label(self, text=self.inv_label).grid(row=4, column=0)
            Label(self, textvariable=self.text_var, bd=10, font="Arial 15 bold",
                  fg="dark slate grey").grid(row=5, column=0)

        self.rows.event_generate("<<ComboboxSelected>>")
        self.pack()

    def col_upd(self,event):
        #Update row-combobox values based on column selection.
        self.column_val = self.columns.get().lower().replace(" ","_")
        if self.column_val == "all":
            self.value_rows = ["N/A"]
        else:
            self.conn = sqlite3.Connection("inventory.db")
            self.conn.row_factory = lambda cursor, row: row[0]
            self.cur = self.conn.cursor()
            if self.column_val == "barrel_no":
                self.cur.execute("SELECT " + self.column_val +
                                  " FROM " + self.sqlite_table)
                self.value_rows = {x[:2] + "-XXX" for x
                                   in self.cur.fetchall()}
            elif "date" in self.column_val: #Get year from date value
                try:
                    self.cur.execute("SELECT " + self.column_val +
                                      " FROM " + self.sqlite_table)
                    self.value_rows = {x[:4] for x in self.cur.fetchall()}
                except sqlite3.OperationalError:
                    self.cur.execute("SELECT po_date" +
                                      " FROM " + self.sqlite_table)
                    self.value_rows = {x[:4] for x in self.cur.fetchall()}
            elif self.column_val == 'pick_up':
                self.cur.execute("SELECT pu_date" +
                                  " FROM " + self.sqlite_table)
                self.value_rows = {x[:4] for x in self.cur.fetchall()}
            elif self.column_val == "age":
                self.cur.execute("SELECT " + self.column_val +
                                  " FROM " + self.sqlite_table)
                self.value_rows = {x[0] + " year(s)" for x
                                   in self.cur.fetchall()}
            elif self.column_val == "product":
                self.cur.execute("SELECT " + self.column_val +
                                  " FROM " + self.sqlite_table)
                self.value_rows = {x for x
                                   in self.cur.fetchall()}
                self.value_rows = list(self.value_rows)
                for ind,item in zip(range(len(self.value_rows)),
                                    self.value_rows):
                    self.mo = re.search("\d",item)
                    #Strip value at location of first digit.
                    if self.mo and self.mo.start() > 1:
                        self.value_rows[ind] = (
                            self.value_rows[ind][0:self.mo.start() - 1])
                    else:
                        continue
                self.value_rows = set(self.value_rows)
            else:
                try:
                    self.cur.execute("SELECT " + self.column_val +
                                      " FROM " + self.sqlite_table)
                    self.value_rows = {x for x
                                       in self.cur.fetchall()}
                except (sqlite3.OperationalError, IndexError):
                    self.value_rows = ["N/A"]
            self.conn.close()
        try:
            self.value_rows = list(self.value_rows)
            self.value_rows.sort()
            self.rows.config(values=self.value_rows)
            self.rows.set(self.value_rows[0])
        except IndexError:
            self.rows.set("N/A")
        self.rows.event_generate("<<ComboboxSelected>>")

    def row_upd(self,event):
        #Update gui_table and total calculation based on row selection.
        self.row_val = self.rows.get()
        if self.row_val == "N/A":
            view_products(self.sqlite_table, "All", "All", self.gui_table)
        else:
            view_products(self.sqlite_table, self.column_val, self.row_val,
                          self.gui_table)
        self.total_calc()


    def total_calc(self):
        #Returns the sum of all values in a table's chosen column.  Used
        #by the View_Frame class to display output.
        self.total = 0
        if self.sqlite_table == "barrels":
            try:
                self.pg_ind = self.gui_table.columns.index("Proof Gallons")
                self.type_ind = self.gui_table.columns.index("Type")
                self.conn = sqlite3.Connection("inventory.db")
                self.conn.row_factory = lambda cursor, row: row[8]
                self.cur = self.conn.cursor()
                self.cur.execute("SELECT * " +
                                   "FROM 'estimated_cogs'")
                self.cogs = self.cur.fetchall()
                self.conn.close()
                for child in self.gui_table.get_children():
                    if (self.gui_table.item(child)["values"][self.type_ind]
                        == "Rum"):
                        self.total += (
                            float(
                              self.gui_table.item(child)["values"][self.pg_ind])
                            * self.cogs[1])
                    else:
                        self.total += (
                            float(
                              self.gui_table.item(child)["values"][self.pg_ind])
                            * self.cogs[0])
            except:
                pass
        else:
            try:
                self.price_ind = self.gui_table.columns.index("Price")
                self.amount_ind = self.gui_table.columns.index("Amount")
                for child in self.gui_table.get_children():
                    self.total += (
                        float(
                          self.gui_table.item(child)["values"][self.price_ind])
                        * float(
                          self.gui_table.item(child)["values"][self.amount_ind])
                        )
            except:
                try:
                    self.price_ind = self.gui_table.columns.index("Price")
                    for child in self.gui_table.get_children():
                        self.total += (
                          float(
                           self.gui_table.item(child)["values"][self.price_ind])
                           )
                except:
                    pass
        try:
            self.text = "{0:,.2f}".format(self.total)
            self.text_var.set("$%s" % (self.text))
        except:
            self.text_var.set("$0.00")

class Barrel_Count_Frame(LabelFrame):
    #Frame containing the number of empty barrels, their price and the
    #total price. Also contains buttons to update these values.
    def __init__(self, master):
        self.master = master
        LabelFrame.__init__(self, master=self.master, text="Barrel Valuation",
                             relief=RIDGE, font="bold", bd=5, padx=2, pady=2)

        (Label(self, text="Full Barrels", bg="dark slate gray", fg="white")
               .grid(row=0, column=0, sticky="NESW"))
        Entry(self, justify='center').grid(row=1, column=0)
        (Label(self, text="Empty Barrels", bg="dark slate gray", fg="white")
               .grid(row=0, column=1, sticky="NESW"))
        Entry(self, justify='center').grid(row=1, column=1)
        (Label(self, text="Price", bg="dark slate gray", fg="white")
               .grid(row=2, column=0, sticky="NESW"))
        Entry(self, justify='center').grid(row=3, column=0)
        (Label(self, text="Total", bg="dark slate gray", fg="white")
               .grid(row=2, column=1, sticky="NESW"))
        Entry(self, justify='center').grid(row=3, column=1)

        # Checkbox used to lock/unlock entries
        self.lock_var = IntVar()
        self.lock_var.set(1)
        self.lock_cbox = Checkbutton(self, text="Lock/Unlock",
                                     variable=self.lock_var,
                                     command=self.cbox_check)
        self.lock_cbox.grid(row=4, column=0)

        (HoverButton(self, text="Update", font='Arial 9 bold',
                     command=self.barr_update)
                     .grid(row=4, column=1, sticky="NESW"))
        self.barr_update(first=1)
        self.pack(anchor='center')

    def cbox_check(self):
        #Checkbox selection function to lock/unlock entries
        if self.lock_var.get() == 1:
            for entry in self.entries[1:3]: #empty barrels / price entry
                entry.config(state='readonly')
        else:
            for entry in self.entries[1:3]: #empty barrels / price entry
                entry.config(state='normal')

    def barr_update(self, first=0):
        self.conn = sqlite3.Connection("inventory.db")
        self.cur = self.conn.cursor()
        self.cur.execute("""SELECT *
                              FROM 'barrel_count'
                         """)
        self.barr_count_vals = list(self.cur.fetchone())
        self.conn.close()
        self.entries = [x for x
                        in reversed(self.grid_slaves())
                        if x.winfo_class() == 'Entry']
        if first == 1:  #First time running function
            self.barr_upd_vals = self.barr_count_vals
        else:
            self.entries_vals = [x.get() for x in self.entries]
            self.sql_edit = self.entries_vals + self.barr_count_vals
            self.sql_edit = tuple(self.sql_edit)
            edit_db(self.sql_edit, 'barrel_count', None, None)
            self.conn = sqlite3.Connection("inventory.db")
            self.cur = self.conn.cursor()
            self.cur.execute("""SELECT *
                                  FROM 'barrel_count'
                             """)
            self.barr_upd_vals = self.cur.fetchone()
            self.conn.close()
        for (entry, val) in zip(self.entries, self.barr_upd_vals):
            entry.config(state='normal')
            entry.delete(0, END)
            entry.insert(0, val)
            entry.config(state='readonly')
            self.lock_var.set(1)


class Cogs_View(Toplevel):
    #Toplevel to view and edit estimated_cogs table.
    def __init__(self, master, sqlite_table, gui_table, view_fr):
        self.master = master
        self.sqlite_table = sqlite_table
        self.gui_table = gui_table
        self.view_fr = view_fr
        self.x = x + 100
        self.y = y + 100
        Toplevel.__init__(self, master=self.master)

        self.whiskey_fr = LabelFrame(self, text="Whiskey COGS",
                                     font="Arial 12 bold")
        self.rum_fr = LabelFrame(self, text="Rum COGS", font="Arial 12 bold")
        self.conn = sqlite3.Connection("inventory.db")
        self.cur = self.conn.cursor()
        self.cur.execute("SELECT * " +
                           "FROM estimated_cogs")
        self.cogs_values = self.cur.fetchall()
        self.conn.close()
        self.cogs_labels = [
            "Raw Materials","Energy","Labor",
            "Error","Total Per Bottle","Bond Ins.",
            "Storage","Mult Factor","Total Per PG"
            ]

        for frame,l in zip([self.whiskey_fr, self.rum_fr],
                           self.cogs_values):
            for ind,item,desc in zip(range(len(l)),
                                     l,
                                     self.cogs_labels):
                Label(frame, text=desc).grid(row=ind, column=0)
                self.ent = Entry(frame, justify="center")
                self.ent.insert(0, item)
                if desc in ["Total Per Bottle", "Total Per PG"]:
                    self.ent.config(state="readonly")
                else:
                    self.ent.config(
                        validate='key',
                        validatecommand=(
                            self.register(valid_dec), '%S', '%s', '%d'))
                self.ent.grid(row=ind, column=1)

        self.whiskey_fr.grid(row=0, column=0, padx=5)
        self.rum_fr.grid(row=0, column=1, padx=5)

        self.button_fr = Frame(self, pady=5)
        (HoverButton(self.button_fr, text="Update", command=self.update)
                     .grid(row=0, column=0, padx=5))
        (HoverButton(self.button_fr, text="Cancel",
                     command=lambda: self.destroy())
                     .grid(row=0, column=1, padx=5))
        self.button_fr.grid(row=1, column=0, columnspan=2)

        self.total_after()
        self.title("COGS")
        self.geometry("%dx%d+%d+%d" % (450, 247, self.x, self.y))
        self.resizable(0,0)
        self.focus()

    def update(self):
        #Update estimated_cogs table with newly inputted values.
        self.conf_ques = messagebox.askquestion("COGS Confirmation",
                            "Are you sure you want to confirm? Confirming will "
                            + "update the Cost of Goods values with your "
                            + "changes.",parent=self)
        if self.conf_ques == "no":
            return self.total_after()
        else:
            self.after_cancel(self.after_func)
            self.whsk_entries = [x.get() for x
                                 in reversed(self.whiskey_fr
                                             .grid_slaves(column=1))]
            self.rum_entries = [x.get() for x
                                in reversed(self.rum_fr.grid_slaves(column=1))]
            self.conn = sqlite3.Connection("inventory.db")
            self.conn.row_factory = lambda cursor, row: row[0]
            self.cur = self.conn.cursor()
            self.cur.execute("SELECT total_per_pg " +
                               "FROM estimated_cogs")
            self.pg_ref = self.cur.fetchall()
            self.whsk_entries.append(self.pg_ref[0])
            self.whsk_entries = tuple(self.whsk_entries)
            self.rum_entries.append(self.pg_ref[1])
            self.rum_entries = tuple(self.rum_entries)
            self.cur.execute("UPDATE estimated_cogs " +
                                "SET raw_mat=?, energy=?, labor=?, error=?, " +
                                     "total_per_bottle=?, bond_ins=?, " +
                                     "storage=?,mult_fact=?, total_per_pg=? " +
                              "WHERE total_per_pg=?",
                              self.whsk_entries)
            self.cur.execute("UPDATE estimated_cogs " +
                                "SET raw_mat=?, energy=?, labor=?, error=?, " +
                                     "total_per_bottle=?, bond_ins=?, " +
                                     "storage=?, mult_fact=?, total_per_pg=? " +
                              "WHERE total_per_pg=?",
                              self.rum_entries)
            self.conn.commit()
            self.conn.close()
            self.view_fr.columns.event_generate("<<ComboboxSelected>>")
            self.destroy()

    def total_after(self):
        #Auto-complete 'total per bottle' and 'total per pg' entries
        #based on other entries.
        def total_update():

            for frame in [self.whiskey_fr, self.rum_fr]:
                self.entries = [x for x
                                in reversed(frame.grid_slaves(column=1))]
                self.entries[4].config(state="normal")
                self.entries[4].delete(0, END)
                self.bot_total = 0
                self.entries[8].config(state="normal")
                self.entries[8].delete(0, END)
                self.pg_total = 0
                for entry in self.entries[:4]:
                    try:
                        self.bot_total += float(entry.get())
                    except:
                        pass
                try:
                    self.entries[4].insert(0, "%.2f" % self.bot_total)
                    self.entries[4].config(state="readonly")
                except:
                    pass
                try:
                    self.pg_total = ((self.bot_total
                                      * float(self.entries[7].get()))
                                     + float(self.entries[5].get())
                                     + float(self.entries[6].get()))
                except:
                    pass
                try:
                    self.entries[8].insert(0, "%.2f" % self.pg_total)
                    self.entries[8].config(state="readonly")
                except:
                    pass
        total_update()
        self.after_func = self.after(150, self.total_after)


class Emptr_View(Toplevel):
    #Toplevel used to input transaction information when inventory items
    #are taken from, or returned to Montgomery.
    def __init__(self, master, sqlite_table, gui_table):
        self.master = master
        self.sqlite_table = sqlite_table
        self.gui_table = gui_table
        self.x = (screen_width/2) - (width/2) + 100
        self.y = ((screen_height/2) - (height/2)) + 50
        self.window_height = 0
        Toplevel.__init__(self, master=self.master)

        self.title_fr = Frame(self)
        (Label(self.title_fr, text="Employee Transaction", font="Arial 10 bold")
              .pack())
        self.title_fr.grid(row=0, column=0, columnspan=2, pady=5)

        self.info_fr = Frame(self)
        for index,desc in enumerate(self.gui_table.columns):
            Label(self.info_fr, text=desc + ":").grid(row=index,column=0)
            if desc == "Date":
                self.date = StringVar()
                self.date_entry = Entry(self.info_fr, state='readonly',
                                        justify='center',
                                        textvariable=self.date)
                self.date_entry.grid(row=index, column=1)
                self.cal_link = HoverButton(self.info_fr, image=cal_photo,
                                            command=lambda:
                                            cal_button(self, self.date_entry))
                self.cal_link.image = cal_photo
                self.cal_link.grid(row=index, column=2)
            elif desc == "Product":
                self.conn = sqlite3.Connection("inventory.db")
                self.conn.row_factory = lambda cursor, row: row[0]
                self.cur = self.conn.cursor()
                self.cur.execute("SELECT product " +
                                   "FROM 'bottles'")
                self.product_rows = self.cur.fetchall()
                self.conn.close()
                self.products = ttk.Combobox(self.info_fr,
                                             values=self.product_rows)
                self.products.config(width=16, background="white",
                                     justify='center', state='readonly')
                self.products.set(self.product_rows[0])
                self.products.grid(row=index, column=1)

            elif desc == "Unit":
                self.units = ttk.Combobox(self.info_fr,
                                          values=['Cases', 'Bottles'])
                self.units.config(width=16, background="white",
                                  justify='center', state='readonly')
                self.units.set('Cases')
                self.units.grid(row=index, column=1)
            elif desc == "Destination":
                self.dest_entry = Entry(self.info_fr)
                self.dest_entry.grid(row=index, column=1)
            else:
                Entry(self.info_fr).grid(row=index,column=1)
            self.window_height += 35
        self.info_fr.grid(row=1,column=0,columnspan=2)
        # Checkbox used indicate Montgomery return
        self.check_var = IntVar()
        self.check_var.set(1)
        self.check_b = Checkbutton(
            self, text="Items were returned to Montgomery",
            variable=self.check_var, command=self.cbox_check)
        self.check_b.grid(row=2, column=0, columnspan=2)
        self.cbox_check()

        self.button_fr = Frame(self)
        (HoverButton(self.button_fr, text="Confirm", width=10,
                     command=lambda: self.confirm())
                     .pack(side=LEFT, padx=5, pady=5))
        (HoverButton(self.button_fr, text="Cancel", width=10,
                     command=lambda: self.destroy())
                     .pack(side=LEFT, padx=5, pady=5))
        self.button_fr.grid(row=3, column=0, columnspan=2)

        self.title("Employee Transasction")
        self.focus()
        self.geometry("+%d+%d" % (self.x, self.y))
        self.resizable(0,0)

    def confirm(self):
        #Update 'employee_transactions' table with entry information.
        #Also update 'bottles'/'samples' inventories.
        self.entries = [x.get() for x
                        in reversed(self.info_fr.grid_slaves())
                        if x.winfo_class() == "Entry"
                        or x.winfo_class() == "TCombobox"]
        if not all(self.entries):
            messagebox.showerror("Input Error",
                                 "Please make sure all of the entries are " +
                                 "completed and then try again.", parent=self)
            return
        if self.entries[3] == "Cases":
            self.inv_tbl = 'bottles'
        else:
            self.inv_tbl = 'samples'
        self.conn = sqlite3.Connection("inventory.db")
        self.cur = self.conn.cursor()
        self.cur.execute("INSERT INTO 'employee_transactions' " +
                              "VALUES (?,?,?,?,?,?)",
                         tuple(self.entries))
        if self.check_var.get() == 1:
            self.cur.execute("UPDATE " + self.inv_tbl +
                               " SET amount=(amount + ?) " +
                              "WHERE product=?",
                             (self.entries[2], self.entries[1]))
        else:
            self.cur.execute("UPDATE " + self.inv_tbl +
                               " SET amount=(amount - ?) " +
                              "WHERE product=?",
                              (self.entries[2], self.entries[1]))
        self.conn.commit()
        self.conn.close()
        db_update()
        view_products('employee_transactions', 'All', 'All', emptr_tbl)
        self.destroy()

    def cbox_check(self):
        #Checkbox selection function
        if self.check_var.get() == 1:
            self.dest_entry.delete(0, END)
            self.dest_entry.insert(0, "Montgomery")
            self.dest_entry.config(state='readonly')
        else:
            self.dest_entry.config(state="normal")
            self.dest_entry.delete(0, END)


class Empty_Barrel_View(Toplevel):
    #Toplevel used to input date for when selected barrel was emptied.
    #Removes barrel from 'barrels' and places in 'empty_barrels'.
    def __init__(self, master, barrel_info):
        self.master = master
        self.barrel_info = barrel_info
        self.x = (screen_width/2) - (width/2) + 100
        self.y = ((screen_height/2) - (height/2)) + 50
        Toplevel.__init__(self, master=self.master)

        self.main_fr = LabelFrame(self,
                                  text="Empty Barrel: " + self.barrel_info[0],
                                  font="Arial 10 bold")
        Label(self.main_fr, text="Empty Date:").grid(row=0, column=0)
        self.date_entry = Entry(self.main_fr, justify='center',
                                state='readonly')
        self.date_entry.grid(row=0, column=1)
        self.cal_link = HoverButton(self.main_fr, image=cal_photo,
                                    command=lambda:
                                    cal_button(self, self.date_entry))
        self.cal_link.image = cal_photo
        self.cal_link.grid(row=0,column=3)
        Label(self.main_fr, text="Remaining PG:").grid(row=1, column=0)
        (Entry(self.main_fr, validate='key',
              validatecommand=(self.register(valid_dec), '%S', '%s', '%d'))
              .grid(row=1, column=1))
        self.add_empty = IntVar()
        self.add_empty.set(1)
        self.add_empty_b = Checkbutton(self.main_fr,
                                       text="Add empty barrel to inventory?",
                                       variable=self.add_empty)
        self.add_empty_b.grid(row=2, column=0, columnspan=2)
        self.main_fr.grid(row=0, column=0, pady=5, padx=10, ipadx=2, ipady=2)

        self.button_fr = Frame(self)
        (HoverButton(self.button_fr, text="Confirm", command=self.confirm)
                     .pack(side=LEFT, padx=5))
        HoverButton(self.button_fr, text="Cancel").pack(side=LEFT, padx=5)
        self.button_fr.grid(row=1, column=0, pady=5)

        self.title("Empty Barrel: " + self.barrel_info[0])
        self.focus()
        self.geometry("+%d+%d" % (self.x, self.y))
        self.resizable(0,0)

    def confirm(self):
        self.conf_quest = messagebox.askquestion(
            "Empty Barrel " + self.barrel_info[0] + "?",
            "Are you sure you want to confirm? Make sure everything is entered "
            + "correctly before continuing.", parent=self)
        if self.conf_quest == "yes":
            self.entry_vals = [x.get() for x
                               in reversed(self.main_fr.grid_slaves())
                               if x.winfo_class() == 'Entry']
            if all(self.entry_vals):
                self.conn = sqlite3.Connection("inventory.db")
                self.cur = self.conn.cursor()
                self.cur.execute("""DELETE FROM barrels
                                          WHERE barrel_no=?""",
                                          (self.barrel_info[0],))
                self.ins_values = (self.barrel_info[:4] +
                                   [self.entry_vals[1]] +
                                   [self.barrel_info[4]] +
                                   [self.entry_vals[0]] +
                                   self.barrel_info[5:])
                self.ins_values = tuple(self.ins_values)
                self.cur.execute("""INSERT INTO empty_barrels
                                         VALUES (?,?,?,?,?,?,?,?,?)""",
                                         (self.ins_values))
                if self.add_empty.get() == 1:
                    self.cur.execute("""UPDATE barrel_count
                                           SET empty_amount=(empty_amount + 1)
                                     """)
                self.conn.commit()
                self.conn.close()
            else:
                messagebox.showerror(
                    "Input Error",
                    "Please make sure all of the entries have values.",
                    parent=self)
        else:
            return
        self.destroy()
        db_update()
        try:
            barr_vfr.columns.set("All")
            barr_vfr.columns.event_generate("<<ComboboxSelected>>")
        except:
            pass
        barr_count_fr.barr_update(first=1)

class Option_Frame(LabelFrame):
    #Frame used to place buttons with inventory functionality.
    def __init__(self, master):
        self.master = master
        self.height = height
        LabelFrame.__init__(self, master=self.master, text="Options",
                            height=self.height, relief=RIDGE, font="bold", bd=5)


class Logistics_Button(HoverButton):
    #Command frame button
    def __init__(self, master, text, sqlite_table, gui_table, command):
        self.sqlite_table = sqlite_table
        self.gui_table = gui_table
        HoverButton.__init__(self, master=master, text=text, width=20, height=1,
                        font=('Calibri', 12, 'bold'), command=command)
        self.pack(anchor='center')


class Treeview_Table(ttk.Treeview):
    #Creates a gui_table with given columns.
    #Has ability to be sorted by gui_table_sort when column headers are
    #clicked upon.
    def __init__(self, master, columns):
        ttk.Treeview.__init__(self, master, columns=columns, show='headings',
                              height=600, style="Custom.Treeview")
        self.width = int(table_width / (len(columns)))
        self.columns = columns
        for i in range(len(columns)):
            self.column(self.columns[i], anchor='center', width=self.width)
            self.heading(str('#' + str((i+1))),
                         text=self.columns[i],
                         command=lambda col=self.columns[i]:
                         gui_table_sort(self, col, False))
        self.pack(side=RIGHT, fill=BOTH, expand=1)


#Used to search for the string literal within a filename that occurs
#before the file extension (Ex. '.txt').
fileRegex = re.compile(r'''
                       ([a-zA-Z0-9_ -]+)
                       (.)
                       ([a-zA-Z_0-9])''',re.VERBOSE)

#Used to search for the the po-number in purchase order file names.
poRegex = re.compile(r'''
                     ([a-zA-Z0-9_]+)
                     (-)
                     ([0-9]{3})
                     (.)
                     ([a-zA-Z_0-9]+)''',re.VERBOSE)

#Used to find important parts of mash number strings.
mashRegex = re.compile(r'''
                       (^\d{4})
                       (/)
                       (\d{2})
                       (-)
                       (\d{1})
                       ([a-zA-Z]{1})''',re.VERBOSE)

def valid_dig(str,act):
    #Entry validation used to ensure only digits.
    if act == '0':
        return True
    else:
        return str.isdigit()


def valid_dec(str,cur_str,act):
    #Entry validation used to ensure only decimal numbers.
    if act == '0':
        return True
    elif str == "." and cur_str.find(".") != -1:
        return False
    elif str =="." and cur_str.find(".") == -1:
        return True
    else:
        return str.isdigit()


def disable_event():
    #Prevent user from 'x'-ing out of sample-desc entry to
    #prevent issue with confirm function.
    pass


#Option values for dropdown menus.
type_options = {
    'raw_materials' : ['Bottles', 'Boxes', 'Caps', 'Capsules', 'Labels'],
    'bottles' : ['Vodka', 'Whiskey', 'Rum', 'Seltzer', 'Other'],
    'barrels' : ['Bourbon', 'Rye', 'Malt', 'Rum', 'Other'],
    'grain' : ['Corn', 'Rye', 'Malted Barley', 'Malted Wheat', 'Wheat', 'Oat',
               'Molasses'],
    'samples' : ['Vodka', 'Whiskey', 'Rum', 'Seltzer', 'Other'],
    'mashes' : ['Bourbon', 'Rye', 'Malt', 'Rum', 'Other'],
    'grain_log' : ['Corn', 'Rye', 'Malted Barley', 'Malted Wheat', 'Oat',
                   'Molasses']}

#Create root window, resize based on user's screen info.
window = Tk()
window.title("Albany Distilling Company Inventory")
window.wm_iconbitmap('favicon.ico')
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()
width = int(screen_width/1.2)
height = int(screen_height/1.2)
command_width = int(.33*width)
table_width = int(.66*width)
x = (screen_width/2) - (width/2)
y = ((screen_height/2) - (height/2)) - 50
window.geometry("%dx%d+%d+%d" % (width, height ,x ,y))
window.resizable(1,1)
window.focus()

#calendar button image
cal_image = Image.open("calendar.png")
cal_image = cal_image.resize((22,22))
cal_photo = ImageTk.PhotoImage(cal_image)

#set styles for gui and certain widgets
s = ttk.Style(window)
s.theme_use('xpnative')
s.configure("Treeview", highlightthickness=0, bd=0, font=('Calibri', 11))
s.element_create("Custom.Treeheading.border", "from", "default")
s.layout("Custom.Treeview.Heading", [
    ("Custom.Treeheading.cell", {'sticky': 'nswe'}),
    ("Custom.Treeheading.border", {'sticky':'nswe', 'children': [
        ("Custom.Treeheading.padding", {'sticky':'nswe', 'children': [
            ("Custom.Treeheading.image", {'side':'right', 'sticky':''}),
            ("Custom.Treeheading.text", {'sticky':'we'})
        ]})
    ]}),
])
s.configure("Custom.Treeview.Heading", background="dark slate grey",
            foreground="white", relief="flat")
s.configure("Treeview.Heading", font=('Calibri', 12,'bold'))
s.configure("TButton", font=('Calibri',12,'bold'))

#Create/load and display database.
inventory = database()

#Create bottle inventory notebook, populate with tabbed frames.
bottinv_nb = ttk.Notebook(window, height=height, width=width)
raw_fr = ttk.Frame(bottinv_nb)
prod_fr = ttk.Frame(bottinv_nb)
inprog_fr = ttk.Frame(bottinv_nb)
bott_fr = ttk.Frame(bottinv_nb)
samp_fr = ttk.Frame(bottinv_nb)

bottinv_nb.add(raw_fr, text="Raw Materials", padding=10)
raw_fr.bind('<Visibility>',
            lambda event:
            raw_vfr.columns.event_generate("<<Comboboxselected>>"))

bottinv_nb.add(prod_fr, text="Production Log", padding=10)
prod_fr.bind('<Visibility>',
             lambda event:
             view_products('production', 'All', 'All', prod_tbl))

bottinv_nb.add(inprog_fr, text="In Progress", padding=10)
inprog_fr.bind('<Visibility>',
               lambda event:
               view_products('in_progress', 'All', 'All', inprog_tbl))

bottinv_nb.add(bott_fr, text="Bottle Inventory", padding=10)
bott_fr.bind('<Visibility>',
             lambda event:
             bott_vfr.columns.event_generate("<<ComboboxSelected>>"))

bottinv_nb.add(samp_fr, text="Samples", padding=10)
samp_fr.bind('<Visibility>',
             lambda event:
             samp_vfr.columns.event_generate("<<ComboboxSelected>>"))

bottinv_nb.pack(side=BOTTOM, fill=BOTH, expand=1)

raw_tbl = Treeview_Table(raw_fr,("Type", "Product", "Amount", "Price", "Total"))
raw_cfr = Command_Frame(raw_fr)
raw_vfr = View_Frame(raw_cfr, 'raw_materials', raw_tbl)
raw_optfr = Option_Frame(raw_cfr)
Logistics_Button(raw_optfr, "Add Item", 'raw_materials', raw_tbl,
                 lambda:
                 Add_View(window, 'raw_materials', raw_tbl, 1, raw_vfr))
Logistics_Button(raw_optfr, "Production", 'raw_materials', raw_tbl,
                 lambda:
                 Production_View(window, 'bottles', raw_tbl))
Logistics_Button(raw_optfr, "Edit Selection", 'raw_materials', raw_tbl,
                 lambda:
                 selection_check('raw_materials', raw_tbl, raw_vfr))
Logistics_Button(raw_optfr, "Delete Selection", 'raw_materials', raw_tbl,
                 lambda:
                 selection_check('raw_materials', raw_tbl, raw_vfr,
                                 delete=True))
raw_optfr.pack()
raw_cfr.pack(padx=10)

prod_tbl = Treeview_Table(prod_fr, ("Date", "Product", "Amount"))
prod_cfr = Command_Frame(prod_fr)
prod_optfr = Option_Frame(prod_cfr)
Logistics_Button(prod_optfr, "Edit Selection", 'production', prod_tbl,
                 lambda:
                 selection_check('production', prod_tbl, None))
Logistics_Button(prod_optfr, "Delete Selection", 'production', prod_tbl,
                 lambda:
                 selection_check('production', prod_tbl, None, delete=True))
prod_optfr.pack()
prod_cfr.pack(padx=10)

inprog_tbl = Treeview_Table(inprog_fr, ("Date", "Product", "Amount",
                                        "Description"))
inprog_cfr = Command_Frame(inprog_fr)
inprog_optfr = Option_Frame(inprog_cfr)
Logistics_Button(inprog_optfr, "Finish Selection", 'in_progress', inprog_tbl,
                 lambda:
                 selection_check(None, inprog_tbl, None, None))
Logistics_Button(inprog_optfr, "Edit Selection", 'in_progress', inprog_tbl,
                 lambda:
                 selection_check('in_progress', inprog_tbl))
Logistics_Button(inprog_optfr, "Delete Selection", 'in_progress', inprog_tbl,
                 lambda:
                 selection_check('in_progress', inprog_tbl, None, delete=True))
inprog_optfr.pack()
inprog_cfr.pack(padx=10)

bott_tbl = Treeview_Table(bott_fr, ("Type", "Product", "Amount", "Case Size",
                                    "Price", "Total"))
bott_cfr = Command_Frame(bott_fr)
bott_vfr = View_Frame(bott_cfr, 'bottles', bott_tbl)
bott_optfr = Option_Frame(bott_cfr)
Logistics_Button(bott_optfr, "Add Item", 'bottles', bott_tbl,
                 lambda:
                 Add_View(window, 'bottles', bott_tbl, 1, bott_vfr))
Logistics_Button(bott_optfr, "Edit Selection", 'bottles', bott_tbl,
                 lambda:
                 selection_check('bottles', bott_tbl, bott_vfr))
Logistics_Button(bott_optfr, "Delete Selection", 'bottles', bott_tbl,
                 lambda:
                 selection_check('bottles', bott_tbl, bott_vfr, delete=True))
bott_optfr.pack()
bott_cfr.pack(padx=10)

samp_tbl = Treeview_Table(samp_fr, ("Type", "Product", "Amount", "Price",
                                    "Total"))
samp_cfr = Command_Frame(samp_fr)
samp_vfr = View_Frame(samp_cfr, 'samples', samp_tbl)
samp_optfr = Option_Frame(samp_cfr)
Logistics_Button(samp_optfr, "Add Item", 'samples', samp_tbl,
                 lambda:
                 Add_View(window, 'samples', samp_tbl, 1, samp_vfr))
Logistics_Button(samp_optfr,"Edit Selection",'samples',samp_tbl,
                 lambda:
                 selection_check('samples', samp_tbl, samp_vfr))
Logistics_Button(samp_optfr, "Delete Selection", 'samples', samp_tbl,
                 lambda:
                 selection_check('samples', samp_tbl, samp_vfr, delete=True))
samp_optfr.pack()
samp_cfr.pack(padx=10)

grain_nb = ttk.Notebook(window, height=height, width=width)
grain_fr = ttk.Frame(grain_nb)
grain_nb.add(grain_fr, text="Grain Inventory", padding=10)
grain_fr.bind('<Visibility>',
              lambda event:
              grain_vfr.columns.event_generate("<<ComboboxSelected>>"))

mash_fr = ttk.Frame(grain_nb)
grain_nb.add(mash_fr, text="Mash Log", padding=10)
mash_fr.bind('<Visibility>',
             lambda event:
             mash_vfr.columns.event_generate("<<ComboboxSelected>>"))

grain_log_fr = ttk.Frame(grain_nb)
grain_nb.add(grain_log_fr, text="Grain Log", padding=10)
grain_log_fr.bind('<Visibility>',
                  lambda event:
                  view_products('grain_log', 'All', 'All', grain_log_tbl))

grain_tbl = Treeview_Table(grain_fr, ("Date", "Order No", "Type", "Amount",
                                      "Price", "Total"))
grain_cfr = Command_Frame(grain_fr)
grain_vfr = View_Frame(grain_cfr, 'grain', grain_tbl)
grain_optfr = Option_Frame(grain_cfr)
Logistics_Button(grain_optfr, "Produce Mash", 'grain', grain_tbl,
                 command=lambda:
                 Mash_Production_View(window, mash_tbl))
Logistics_Button(grain_optfr, "Mash Production Sheet", 'grain', grain_tbl, None)
Logistics_Button(grain_optfr, "Add Grain", 'grain', grain_tbl,
                 lambda:
                 Add_View(window, 'grain', grain_tbl, 1, grain_vfr))
Logistics_Button(grain_optfr, "Edit Selection", 'grain', grain_tbl,
                 lambda:
                 selection_check('grain', grain_tbl, grain_vfr))
Logistics_Button(grain_optfr, "Delete Selection", 'grain', grain_tbl,
                 lambda:
                 selection_check('grain', grain_tbl, grain_vfr, delete=True))
grain_optfr.pack()
grain_cfr.pack(padx=10)

mash_tbl = Treeview_Table(mash_fr, ("Date", "Mash No", "Type", "Grains"))
mash_cfr = Command_Frame(mash_fr)
mash_vfr = View_Frame(mash_cfr, 'mashes', mash_tbl)
mash_optfr = Option_Frame(mash_cfr)
Logistics_Button(mash_optfr, "Add Mash", 'mashes', mash_tbl,
                 lambda:
                 Add_View(window, "mashes", mash_tbl, 1, mash_vfr))
Logistics_Button(mash_optfr, "Edit Selection", 'mashes', mash_tbl,
                 lambda:
                 selection_check('mashes', mash_tbl, mash_vfr))
Logistics_Button(mash_optfr, "Delete Selection", 'mashes', mash_tbl,
                 lambda:
                 selection_check('mashes', mash_tbl, mash_vfr, delete=True))
mash_optfr.pack()
mash_cfr.pack(padx=10)

grain_log_tbl = Treeview_Table(grain_log_fr, ("Arrival Date", "Finish Date",
                                              "Type", "Order No"))
grain_log_cfr = Command_Frame(grain_log_fr)
grain_log_optfr = Option_Frame(grain_log_cfr)
Logistics_Button(grain_log_optfr, "Edit Selection", 'grain_log', grain_log_tbl,
                 lambda:
                 selection_check('grain_log', grain_log_tbl, None))
Logistics_Button(grain_log_optfr, "Delete Selection", 'grain_log',
                 grain_log_tbl,
                 lambda:
                 selection_check('grain_log', grain_log_tbl, None, delete=True))
grain_log_optfr.pack()
grain_log_cfr.pack(padx=10)

barr_nb = ttk.Notebook(window, height=height, width=width)
barr_fr = ttk.Frame(barr_nb)
empt_barr_fr = ttk.Frame(barr_nb)

barr_nb.add(barr_fr, text="Barrel Inventory", padding=10)
barr_fr.bind('<Visibility>',
              lambda event:
              barr_vfr.columns.event_generate("<<ComboboxSelected>>"))

barr_nb.add(empt_barr_fr, text="Emptied Barrels", padding=10)
empt_barr_fr.bind('<Visibility>',
                  lambda event:
                  view_products('empty_barrels', 'All', 'All', empt_barr_tbl))

barr_tbl = Treeview_Table(barr_fr, ("Barrel No","Type","Gallons",
                                    "Proof Gallons", "Date Filled", "Age",
                                    "Investor"))
barr_cfr = Command_Frame(barr_fr)
barr_vfr = View_Frame(barr_cfr, 'barrels', barr_tbl)
barr_optfr = Option_Frame(barr_cfr)
Logistics_Button(barr_optfr, "Add Barrel", 'barrels', barr_tbl,
                 lambda:
                 Add_View(window, 'barrels', barr_tbl, 1, barr_vfr))
Logistics_Button(barr_optfr, "Empty Barrel", 'barrels', barr_tbl,
                 lambda:
                 selection_check('barrels', barr_tbl, barr_vfr, empty=True))
Logistics_Button(barr_optfr, "Update COGS", 'barrels', barr_tbl,
                 lambda:
                 Cogs_View(window, 'estimated_cogs', barr_tbl, barr_vfr))
Logistics_Button(barr_optfr, "Edit Selection", 'barrels', barr_tbl,
                 lambda:
                 selection_check('barrels', barr_tbl, barr_vfr))
Logistics_Button(barr_optfr, "Delete Selection", 'barrels', barr_tbl,
                 lambda:
                 selection_check('barrels', barr_tbl, barr_vfr, delete=True))
barr_optfr.pack(pady=2)
barr_count_fr = Barrel_Count_Frame(barr_cfr)
barr_cfr.pack(padx=10)

empt_barr_tbl = Treeview_Table(empt_barr_fr, ('Barrel No', 'Type', 'Gallons',
                                              'PG', 'PG Leftover',
                                              'Filled', 'Emptied',
                                              'Age', 'Investor'))

empt_barr_cfr = Command_Frame(empt_barr_fr)
empt_barr_optfr = Option_Frame(empt_barr_cfr)
Logistics_Button(empt_barr_optfr, "Edit Selection", 'empty_barrels',
                 empt_barr_tbl,
                 lambda:
                 selection_check('empty_barrels', empt_barr_tbl, None))
Logistics_Button(empt_barr_optfr, "Delete Selection", 'empty_barrels',
                 empt_barr_tbl,
                 lambda:
                 selection_check('empty_barrels', empt_barr_tbl, None,
                                 delete=True))
empt_barr_optfr.pack(pady=2)
empt_barr_cfr.pack(padx=10)

po_nb = ttk.Notebook(window, height=height, width=width)
po_fr = Frame(po_nb)
po_nb.add(po_fr, text="Purchase Orders", padding=10)
po_fr.bind('<Visibility>',
              lambda event:
              po_vfr.columns.event_generate("<<ComboboxSelected>>"))

pending_fr = Frame(po_nb)
po_nb.add(pending_fr, text="Pending Purchase Orders", padding=10)
pending_fr.bind('<Visibility>',
              lambda event:
              pending_vfr.columns.event_generate("<<ComboboxSelected>>"))

po_tbl = Treeview_Table(po_fr, ("Date", "Pick Up", "Product", "Amount", "Unit",
                                "Price", "Total", "Destination", "PO No."))
po_cfr = Command_Frame(po_fr)
po_vfr = View_Frame(po_cfr, 'purchase_orders', po_tbl)
po_optfr = Option_Frame(po_cfr)
Logistics_Button(po_optfr, "Create Purchase Order", 'purchase_orders', po_tbl,
                 lambda:
                 Purchase_Order_View(window))
Logistics_Button(po_optfr, "View Purchase Order", 'purchase_orders', po_tbl,
                 lambda:
                 selection_check(None, po_tbl, None, edit=False))
Logistics_Button(po_optfr, "Edit Selection", 'purchase_orders', po_tbl,
                 lambda:
                 selection_check('purchase_orders', po_tbl, po_vfr))
Logistics_Button(po_optfr, "Delete Selection", 'purchase_orders', po_tbl,
                 lambda:
                 selection_check('purchase_orders', po_tbl, po_vfr,
                                 delete=True))
po_optfr.pack()
po_cfr.pack(padx=10)

pending_tbl = Treeview_Table(pending_fr, ("Date", "Pick Up", "Product",
                                          "Amount", "Unit", "Price", "Total",
                                          "Destination", "PO No."))
pending_cfr = Command_Frame(pending_fr)
pending_vfr = View_Frame(pending_cfr, 'pending_po', pending_tbl)
pending_optfr = Option_Frame(pending_cfr)
Logistics_Button(pending_optfr, "Fulfill Purchase Order", 'pending_po',
                 pending_tbl,
                 lambda:
                 selection_check('pending_po', pending_tbl, pending_vfr, False))
Logistics_Button(pending_optfr, "Edit Selection", 'pending_po', pending_tbl,
                 lambda:
                 selection_check('pending_po', pending_tbl, pending_vfr))
Logistics_Button(pending_optfr, "Delete Selection", 'pending_po', pending_tbl,
                 lambda:
                 selection_check('pending_po', pending_tbl, pending_vfr,
                                 delete=True))
pending_optfr.pack()
pending_cfr.pack(padx=10)

emptr_nb = ttk.Notebook(window, height=height, width=width)
emptr_fr = Frame(emptr_nb)
emptr_nb.add(emptr_fr, text="Employee Transactions", padding=10)

emptr_tbl = Treeview_Table(emptr_fr, ("Date", "Product", "Amount", "Unit",
                                      "Employee", "Destination"))
emptr_cfr = Command_Frame(emptr_fr)
emptr_optfr = Option_Frame(emptr_cfr)
Logistics_Button(emptr_optfr, "Transaction", 'employee_transactions',
                 emptr_tbl,
                 lambda:
                 Emptr_View(window, 'employee_transactions', emptr_tbl))
Logistics_Button(emptr_optfr, "Edit Selection", 'employee_transactions',
                 emptr_tbl,
                 lambda:
                 selection_check('employee_transactions', emptr_tbl,
                                 None))
Logistics_Button(emptr_optfr, "Delete Selection", 'employee_transactions',
                 emptr_tbl,
                 lambda:
                 selection_check('employee_transactions', emptr_tbl, None,
                                 delete=True))
emptr_optfr.pack()
emptr_cfr.pack(padx=10)

reports_nb = ttk.Notebook(window, height=height, width=width)
reports_fr = Reports_Frame(reports_nb)
reports_nb.add(reports_fr, text="Monthly Report", padding=10)
reports_fr.bind("<Visibility>",
                lambda event:
                reports_fr.year_cmbo_box.event_generate("<<ComboboxSelected>>"))

menubar = Menu(window)
menu1 = Menu(menubar, tearoff=0)
menu1.add_command(label="Raw Materials and Bottles",
                  command=lambda:
                  view_widget(window, bottinv_nb, BOTTOM, 'raw_materials',
                              'All', 'All', raw_tbl))
menu1.add_command(label="Grain",
                  command=lambda:
                  view_widget(window, grain_nb, BOTTOM, 'grain', 'All', 'All',
                              grain_tbl))
menu1.add_command(label="Barrels",
                  command=lambda:
                  view_widget(window, barr_nb, BOTTOM, 'barrels', 'All', 'All',
                              barr_tbl))
menubar.add_cascade(label="Inventory", menu=menu1)

menu2 = Menu(menubar, tearoff=0)
menu2.add_command(label="Purchase Orders",
                  command=lambda:
                  view_widget(window, po_nb, BOTTOM, 'purchase_orders', 'All',
                              'All', po_tbl))
menu2.add_command(label="Employee Transactions",
                  command=lambda:
                  view_widget(window, emptr_nb, BOTTOM, 'employee_transactions',
                              'All', 'All', emptr_tbl))
menubar.add_cascade(label="Shipping and Transactions", menu=menu2)

menu3 = Menu(menubar, tearoff=0)
menu3.add_command(label="Production Sheets",
                  command=lambda: file_view("production_sheets"))
menu3.add_command(label="Case Labels", command=lambda: file_view("case_labels"))
menubar.add_cascade(label="Files", menu=menu3)

menu4 = Menu(menubar, tearoff=0)
menu4.add_command(label="Monthly Reports",
                  command=lambda: view_widget(window, reports_nb, BOTTOM, None,
                                              'All', 'All', None))
menu4.add_command(label="Export for Excel", command=create_excel_inv)
menubar.add_cascade(label="Analysis", menu=menu4)

window.config(menu=menubar)
window.mainloop()
